import asyncio
import base64
import hashlib
import json
import mimetypes
import os
import re
import subprocess
import tempfile
import time
import traceback
from contextlib import asynccontextmanager
from importlib.metadata import PackageNotFoundError, version
from pathlib import Path
from typing import Any, Optional
from urllib.parse import unquote, urlparse

import httpx
from anyio import to_thread
from docx import Document as WordDocument
from docx.table import Table as WordTable
from docx.text.paragraph import Paragraph as WordParagraph
from fastapi import FastAPI, HTTPException
from fastapi.responses import Response
from fastapi.encoders import jsonable_encoder

from app.config import (
    ACTIVE_BACKENDS,
    DEBUG_ERRORS,
    DOCLING_INSTALLED,
    EXTRACTION_BACKEND,
    FILE_DOWNLOAD_TIMEOUT_SECONDS,
    GENERIC_MIME_TYPES,
    HEAVY_WORK_CONCURRENCY,
    RENDER_PDF_DPI,
    RENDER_PDF_JPEG_QUALITY,
    GEOMETRY_ENRICHMENT_ENABLED,
    IMAGE_SUFFIXES,
    LLAMAPARSE_API_KEY,
    LLAMAPARSE_BASE_URL,
    LLAMAPARSE_LANGUAGE,
    LLAMAPARSE_MAX_RETRIES,
    LLAMAPARSE_MAX_WAIT_SECONDS,
    LLAMAPARSE_POLLING_INTERVAL,
    LLAMAPARSE_REQUEST_TIMEOUT_SECONDS,
    LLAMAPARSE_RESULT_TYPE,
    LLM_IS_YANDEX,
    OPENROUTER_API_KEY,
    OPENROUTER_APP_NAME,
    OPENROUTER_BASE_URL,
    OPENROUTER_MAX_TOKENS,
    OPENROUTER_MODEL,
    OPENROUTER_PDF_ENGINE,
    OPENROUTER_PROVIDER_IGNORE,
    OPENROUTER_PROVIDER_ORDER,
    OPENROUTER_SITE_URL,
    PYMUPDF_INSTALLED,
    REMOTE_API_TIMEOUT_SECONDS,
    STUBBED_BACKENDS,
    SUPPORTED_BACKENDS,
    fitz,
    logger,
)
from app.models import DownloadedFile, ExtractionRequest, PdfPageIndex, PdfWord
from app.schema_utils import normalize_json_schema


_strict_schema_supported: dict[str, bool] = {}


@asynccontextmanager
async def lifespan(app: FastAPI):
    app.state.docling_extractor = None
    try:
        to_thread.current_default_thread_limiter().total_tokens = HEAVY_WORK_CONCURRENCY
        logger.info("Heavy-work concurrency limited to %s", HEAVY_WORK_CONCURRENCY)
    except Exception:
        logger.warning("Failed to set thread limiter", exc_info=True)
    logger.info("Extraction service started; default backend=%s", EXTRACTION_BACKEND)
    try:
        yield
    finally:
        extractor = getattr(app.state, "docling_extractor", None)
        close = getattr(extractor, "close", None)
        if callable(close):
            try:
                close()
            except Exception:
                logger.exception("Failed to close Docling extractor")


app = FastAPI(lifespan=lifespan, title="extraction-service", version="0.3.0")


@app.get("/health")
async def health() -> dict:
    return {
        "status": "ok",
        "default_backend": EXTRACTION_BACKEND,
        "active_backends": sorted(ACTIVE_BACKENDS),
        "stubbed_backends": sorted(STUBBED_BACKENDS),
        "docling_installed": DOCLING_INSTALLED,
        "pymupdf_installed": PYMUPDF_INSTALLED,
        "geometry_enrichment_enabled": GEOMETRY_ENRICHMENT_ENABLED,
        "llamaparse_configured": bool(LLAMAPARSE_API_KEY),
    }


def _get_docling_version() -> Optional[str]:
    if not DOCLING_INSTALLED:
        return None
    try:
        return version("docling")
    except PackageNotFoundError:
        return None


def _select_backend(requested_backend: Optional[str]) -> str:
    backend = (requested_backend or EXTRACTION_BACKEND).strip().lower()
    if backend not in SUPPORTED_BACKENDS:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported backend: {requested_backend or backend}",
        )
    return backend


def _raise_docling_backend_unavailable(backend: str) -> None:
    raise HTTPException(
        status_code=501,
        detail=(
            f"Backend '{backend}' is temporarily disabled in this build. "
            "Use 'openrouter'. Docling integration is kept as a stub for future re-enable."
        ),
    )


def _looks_like_image(filename: str, content_type: str | None) -> bool:
    if content_type and content_type.startswith("image/"):
        return True
    return Path(filename).suffix.lower() in IMAGE_SUFFIXES


def _looks_like_docx(filename: str, content_type: str | None) -> bool:
    normalized = (content_type or "").split(";", 1)[0].strip().lower()
    suffix = Path(filename).suffix.lower()
    return (
        normalized
        in {
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        }
        or suffix in {".docx", ".doc"}
    )


def _looks_like_excel(filename: str, content_type: str | None) -> bool:
    normalized = (content_type or "").split(";", 1)[0].strip().lower()
    suffix = Path(filename).suffix.lower()
    return (
        normalized
        in {
            "application/vnd.ms-excel",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        }
        or suffix in {".xls", ".xlsx", ".xlsm"}
    )


def _looks_like_pdf(filename: str, content_type: str | None) -> bool:
    normalized = (content_type or "").split(";", 1)[0].strip().lower()
    suffix = Path(filename).suffix.lower()
    return normalized == "application/pdf" or suffix == ".pdf"


# Таймаут на ОДИН вызов Tesseract (на страницу). Нормальный OCR страницы — секунды;
# но на некоторых сканах с шумом/паттернами Tesseract LSTM зацикливается и висит
# минутами, блокируя всю задачу извлечения навсегда. Лимит превращает зависание в
# обрабатываемую ошибку (страница пропускается / падаем на другой путь).
OCR_PAGE_TIMEOUT_SECONDS = 45
OCR_OSD_TIMEOUT_SECONDS = 20


def _auto_orient_for_ocr(img):
    """Доворачивает изображение страницы в правильную ориентацию перед OCR.

    Некоторые PDF имеют повёрнутые страницы (rotation 90/270): после рендера в
    картинку текст оказывается боком/вверх ногами, и Tesseract выдаёт мусор.
    Через OSD определяем угол и доворачиваем. При ошибке/низкой уверенности —
    возвращаем как есть (не хуже прежнего поведения)."""
    try:
        import pytesseract

        osd = pytesseract.image_to_osd(
            img, output_type=pytesseract.Output.DICT, timeout=OCR_OSD_TIMEOUT_SECONDS
        )
        rotate = int(osd.get("rotate", 0) or 0)
        conf = float(osd.get("orientation_conf", 0) or 0)
        if rotate and conf >= 1.0:
            # PIL.rotate крутит против часовой; OSD 'rotate' — на сколько повернуть
            # ПО часовой, чтобы выпрямить → используем expand и отрицательный угол.
            return img.rotate(-rotate, expand=True)
    except Exception:
        pass
    return img


def _text_layer_is_usable(text: str) -> bool:
    """Проверяет, что извлечённый текстовый слой ОСМЫСЛЕННЫЙ, а не «мусор».

    У некоторых PDF шрифт со сломанной кодировкой (нет/битый ToUnicode): визуально
    документ читается, но get_text() возвращает мешанину одиночных символов вроде
    'u E e 5 ^s * 2 FE'. Такой текст нельзя слать в LLM — он не найдёт ни одной
    характеристики. Признак мусора: мало «слов» (последовательностей букв >=3).
    Тогда лучше упасть на OCR (он читает отрисованные глифы, а не битую кодировку).
    """
    stripped = text.strip()
    if len(stripped) < 200:
        # Слишком мало текста, чтобы судить о качестве — не блокируем (OCR решит по длине).
        return len(stripped) > 20
    letters = re.findall(r"[^\W\d_]", stripped, re.UNICODE)
    if not letters:
        return False
    words = re.findall(r"[^\W\d_]{3,}", stripped, re.UNICODE)
    letters_in_words = sum(len(w) for w in words)
    word_ratio = letters_in_words / len(letters)
    # У нормального текста (RU/EN) word_ratio ~0.9; у мусора одиночных символов ~0.4-0.5.
    if word_ratio < 0.6:
        return False

    # Битая кодировка кириллического шрифта (нет/сломан ToUnicode): get_text() даёт
    # «слова», визуально читаемые как русский, но это мусор. word_ratio при этом
    # высокий, и такой текст проходил проверку выше. Ловим двумя признаками:
    if len(words) >= 20:
        def _is_cyr(ch: str) -> bool:
            return "Ѐ" <= ch <= "ӿ"

        def _is_lat(ch: str) -> bool:
            return ("a" <= ch <= "z") or ("A" <= ch <= "Z")

        cyrillic = sum(1 for ch in letters if _is_cyr(ch))
        cyrillic_ratio = cyrillic / len(letters)

        mixed_case = 0      # 'HanuenoaaHne' — заглавная не в начале слова
        mixed_script = 0    # 'ВзiпIеll' — кириллица И латиница в одном слове
        for w in words:
            inner = w[1:]
            if any(ch.isupper() for ch in inner) and any(ch.islower() for ch in inner):
                mixed_case += 1
            has_cyr = any(_is_cyr(ch) for ch in w)
            has_lat = any(_is_lat(ch) for ch in w)
            if has_cyr and has_lat:
                mixed_script += 1
        mixed_case_ratio = mixed_case / len(words)
        mixed_script_ratio = mixed_script / len(words)

        # 1. Латиница-вместо-кириллицы ('HanuenoaaHne'): кириллицы почти нет,
        #    но много «рваного» регистра.
        if cyrillic_ratio < 0.15 and mixed_case_ratio >= 0.2:
            return False
        # 2. Кириллица + латиница намешаны в одних словах ('ВзiпIеll', 'llсilrr'):
        #    в нормальном тексте смешение скриптов внутри слова — единичные случаи
        #    (бренды вроде 'iPhone'), у битой кодировки — массово.
        if mixed_script_ratio >= 0.15:
            return False

    return True


def _text_layer_has_enough_content(all_pages_text: list[str]) -> bool:
    """Проверяет, что текстового слоя ДОСТАТОЧНО для извлечения, а не несколько
    символов на скан-страницах (типичный даташит-картинка с одним артикулом в
    углу как текст, остальное — изображение). Без этой проверки система считает
    «150TBS11 × 3» годным слоем и НЕ запускает OCR → LLM получает пустоту.

    Критерий: уникального текста должно быть достаточно относительно числа
    непустых страниц. Учитываем дублирование (на каждой странице один артикул)."""
    non_empty = [t.strip() for t in all_pages_text if t and t.strip()]
    if not non_empty:
        return False
    # Уникальные строки контента (убираем повторяющиеся колонтитулы/артикулы)
    unique_lines: set[str] = set()
    for page_text in non_empty:
        for line in page_text.splitlines():
            line = line.strip()
            if len(line) >= 3:
                unique_lines.add(line)
    unique_chars = sum(len(line) for line in unique_lines)
    page_count = len(all_pages_text) or 1
    # Минимум ~60 уникальных символов на страницу — иначе это скан без текста.
    return unique_chars >= 60 * page_count and unique_chars >= 120


def _page_words_are_usable(pages: list["PdfPageIndex"]) -> bool:
    """Проверяет, что слова в текстовом индексе страниц — осмысленные, а не мусор.

    Используется геометрией: текстовый слой может «существовать» (слова есть), но
    из-за битой кодировки шрифта это мешанина одиночных символов. По такому индексу
    цитаты не находятся → 0 координат → клик по характеристике ничего не делает.
    Если слова мусорные — геометрия должна перейти на OCR-индекс."""
    sample = " ".join(
        word.text
        for page in pages
        for word in page.words
    )
    return _text_layer_is_usable(sample)


def _page_words_are_enough(pages: list["PdfPageIndex"]) -> bool:
    """Достаточно ли уникальных слов в индексе. Скан-PDF с одним артикулом-текстом
    на странице (остальное — картинка) даёт пару слов на лист — этого мало для
    привязки цитат, нужен OCR-индекс."""
    unique_words = {
        word.normalized
        for page in pages
        for word in page.words
        if len(word.normalized) >= 2
    }
    page_count = len(pages) or 1
    return len(unique_words) >= 8 * page_count and len(unique_words) >= 15


def _extract_target_names_from_prompt(prompt: str) -> list[str] | None:
    """Extracts target characteristic names from the prompt JSON list, if present."""
    match = re.search(r"characteristic names to return.*?\n\s*\[", prompt, re.DOTALL | re.IGNORECASE)
    if not match:
        return None
    start = prompt.index("[", match.start())
    try:
        names = json.loads(prompt[start:prompt.index("]", start) + 1])
        if isinstance(names, list):
            return [str(n) for n in names if isinstance(n, str) and n.strip()]
    except (json.JSONDecodeError, ValueError):
        pass
    return None


_RELEVANCE_KEYWORDS = re.compile(
    r"характерист|параметр|показател|спецификац|specification|parameter|"
    r"техничес|nominal|dimension|габарит|размер|масс[аы]|вес\b|"
    r"мощност|напор|подач[аи]|производител|давлен|температур|"
    r"частот|оборот|диаметр|длин[аы]|ширин[аы]|высот[аы]|"
    r"марк[аи]|модел[ьи]|тип\b|обозначен",
    re.IGNORECASE,
)


def _page_is_relevant(text: str, target_names: list[str] | None = None) -> bool:
    if _RELEVANCE_KEYWORDS.search(text):
        return True
    if target_names:
        text_lower = text.lower().replace("ё", "е")
        for name in target_names:
            name_words = [w for w in re.split(r"[\s:,]+", name.lower().replace("ё", "е")) if len(w) >= 3]
            if name_words and sum(1 for w in name_words if w in text_lower) >= max(1, len(name_words) // 2):
                return True
    return False


def _page_has_table(text: str) -> bool:
    """Страница содержит Markdown-таблицу (вставленную _page_text_with_tables):
    есть строка-разделитель «| --- | ...». Самый дешёвый и точный признак."""
    return bool(re.search(r"^\s*\|[\s|:-]*---[\s|:-]*\|", text, re.MULTILINE))


def _page_continues_table(text: str) -> bool:
    """Страница ВЫГЛЯДИТ как продолжение таблицы: начинается со строк «| ... |»,
    но без строки-заголовка-разделителя «| --- |». Так бывает, когда таблица
    переходит на следующую страницу — шапка осталась на предыдущей."""
    stripped = text.lstrip()
    starts_with_row = stripped.startswith("|")
    return starts_with_row and not _page_has_table(text)


def _filter_relevant_pages(
    all_pages_text: list[str],
    target_names: list[str] | None = None,
    max_chars: int = 30000,
) -> tuple[list[tuple[int, str]], bool]:
    """Returns ([(page_number, text), ...], was_filtered).
    Keeps first page, last page, and any page with characteristic-related keywords.
    Falls back to all pages if total text is small enough.

    Пункт 4 — НЕ разрывать таблицы фильтром: страница с таблицей всегда релевантна,
    а её соседи (предыдущая/следующая) подтягиваются целиком, чтобы шапку таблицы
    не оторвало от строк-продолжения на соседней странице."""
    total_len = sum(len(t) for t in all_pages_text)
    if total_len <= max_chars:
        return [(i + 1, t) for i, t in enumerate(all_pages_text) if t.strip()], False

    n = len(all_pages_text)
    keep: set[int] = set()
    for i, text in enumerate(all_pages_text):
        if not text.strip():
            continue
        if i == 0 or i == n - 1:
            keep.add(i)
            continue
        if _page_has_table(text) or _page_is_relevant(text, target_names):
            keep.add(i)

    # Целостность таблиц: для каждой страницы с таблицей или её продолжением
    # подтягиваем непустых соседей, чтобы заголовок и строки не разорвались.
    for i, text in enumerate(all_pages_text):
        if not text.strip():
            continue
        if _page_has_table(text) or _page_continues_table(text):
            for j in (i - 1, i + 1):
                if 0 <= j < n and all_pages_text[j].strip():
                    keep.add(j)

    relevant = [(i + 1, all_pages_text[i]) for i in sorted(keep)]

    if not relevant:
        return [(i + 1, t) for i, t in enumerate(all_pages_text) if t.strip()], False

    # If filtering cut too aggressively (< 3 pages), fall back to full text
    if len(relevant) < 3 and n > 5:
        return [(i + 1, t) for i, t in enumerate(all_pages_text) if t.strip()], False

    return relevant, True


def _augment_sparse_pages_with_ocr(
    document: Any,
    all_pages_text: list[str],
    *,
    sparse_threshold: int = 150,
    min_ocr_gain: int = 200,
) -> tuple[list[str], int]:
    """Для страниц с бедным текстовым слоем прогоняет OCR и подменяет их текст.

    Кейс: документ в основном текстовый (описание читается), но таблицы ТТХ
    нарисованы векторной графикой/картинкой без ToUnicode — get_text() для такой
    страницы почти пуст, и характеристики теряются. Если OCR такой страницы даёт
    существенно больше текста (min_ocr_gain), берём OCR-вариант.

    Возвращает (обновлённый список текстов, число заменённых страниц)."""
    sparse_indices = [
        i for i, t in enumerate(all_pages_text)
        if len((t or "").strip()) < sparse_threshold
    ]
    if not sparse_indices:
        return all_pages_text, 0
    try:
        import pytesseract
        from PIL import Image
        import io
        pytesseract.get_tesseract_version()
    except Exception:
        return all_pages_text, 0

    DPI = 200
    updated = list(all_pages_text)
    replaced = 0
    for i in sparse_indices:
        try:
            page = document.load_page(i)
            pix = page.get_pixmap(matrix=fitz.Matrix(DPI / 72, DPI / 72), alpha=False)
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            img = _auto_orient_for_ocr(img)
            ocr_text = pytesseract.image_to_string(
                img, lang="rus+eng", timeout=OCR_PAGE_TIMEOUT_SECONDS
            ).strip()
        except Exception:
            continue
        if len(ocr_text) >= len((updated[i] or "").strip()) + min_ocr_gain:
            updated[i] = ocr_text
            replaced += 1
    return updated, replaced


def _serialize_table_markdown(rows: list[list[Any]]) -> str:
    """Сериализует таблицу (список строк-списков от find_tables().extract()) в
    Markdown. Первая строка — заголовок. Пустые ячейки (None) → ''; переводы строк
    внутри ячеек → пробел. Markdown даёт модели явную привязку
    «значение → столбец», убирая угадывание по пробелам в плоском тексте."""
    def _cell(value: Any) -> str:
        text = "" if value is None else str(value)
        return re.sub(r"\s+", " ", text).replace("|", "/").strip()

    norm_rows = [[_cell(c) for c in row] for row in rows if row]
    if not norm_rows:
        return ""
    width = max(len(r) for r in norm_rows)
    norm_rows = [r + [""] * (width - len(r)) for r in norm_rows]
    header = norm_rows[0]
    out = ["| " + " | ".join(header) + " |", "| " + " | ".join(["---"] * width) + " |"]
    for row in norm_rows[1:]:
        out.append("| " + " | ".join(row) + " |")
    return "\n".join(out)


def _page_text_with_tables(page: Any) -> str:
    """Текст страницы, где таблицы заменены Markdown-представлением.

    Плоский get_text() в таблицах склеивает значения пробелами, заголовки
    оторваны от строк — LLM угадывает «модель→колонка→значение» и ошибается на
    широких таблицах. find_tables() распознаёт сетку, и мы подаём её как Markdown.
    Нетабличный текст сохраняется в исходном вертикальном порядке. При любой
    ошибке — фолбэк на плоский sorted-текст (надёжность важнее)."""
    try:
        tables_finder = page.find_tables()
        tables = list(tables_finder.tables)
    except Exception:
        return page.get_text("text", sort=True)

    if not tables:
        return page.get_text("text", sort=True)

    table_bboxes = []
    table_segments = []  # (y0, markdown)
    for table in tables:
        try:
            rows = table.extract()
            md = _serialize_table_markdown(rows)
            bbox = fitz.Rect(table.bbox)
        except Exception:
            continue
        if md:
            table_bboxes.append(bbox)
            table_segments.append((float(bbox.y0), md))

    if not table_segments:
        return page.get_text("text", sort=True)

    # Текстовые блоки вне областей таблиц (чтобы не дублировать табличный текст).
    text_segments = []  # (y0, text)
    try:
        blocks = page.get_text("blocks", sort=True)
    except Exception:
        blocks = []
    for block in blocks:
        if len(block) < 5:
            continue
        x0, y0, x1, y1, btext = block[0], block[1], block[2], block[3], block[4]
        if not isinstance(btext, str) or not btext.strip():
            continue
        block_rect = fitz.Rect(x0, y0, x1, y1)
        # Пропускаем блоки, существенно пересекающиеся с любой таблицей.
        inside_table = False
        for tb in table_bboxes:
            inter = block_rect & tb
            if inter.is_valid and inter.get_area() > 0.5 * block_rect.get_area():
                inside_table = True
                break
        if not inside_table:
            text_segments.append((float(y0), btext.strip()))

    # Собираем в вертикальном порядке: текстовые блоки и таблицы вперемешку по Y.
    combined = sorted(text_segments + table_segments, key=lambda s: s[0])
    parts = [seg[1] for seg in combined]
    result = "\n".join(parts).strip()
    return result or page.get_text("text", sort=True)


def _convert_pdf_to_structured_text(
    local_path: str,
    target_names: list[str] | None = None,
) -> dict[str, Any]:
    """Извлекает текст из PDF для передачи в LLM.

    Для больших документов (>30K символов) фильтрует только страницы с
    релевантным содержимым (таблицы характеристик), чтобы сократить prompt.
    """
    if not PYMUPDF_INSTALLED:
        return {"text": "", "page_count": 0, "ocr_applied": False, "error": "PyMuPDF not installed"}

    document = fitz.open(local_path)
    page_count = document.page_count
    lines: list[str] = []
    ocr_applied = False

    all_pages_text: list[str] = []
    for i in range(page_count):
        page = document.load_page(i)
        # Таблицы подаём как Markdown (явная привязка значение→столбец), остальной
        # текст — как обычно. На сканах/OCR ветка ниже не затрагивается.
        text = _page_text_with_tables(page).strip()
        all_pages_text.append(text)

    combined_text = "\n".join(t for t in all_pages_text if t)
    has_enough = _text_layer_has_enough_content(all_pages_text)
    has_text = (
        len(combined_text) > 20
        and _text_layer_is_usable(combined_text)
        and has_enough
    )
    if len(combined_text) > 20 and not has_text:
        if not has_enough:
            logger.info(
                "PDF text layer too sparse (%d chars over %d pages — likely scanned) — falling back to OCR",
                len(combined_text), page_count,
            )
        else:
            logger.info("PDF text layer looks garbled (broken font encoding) — falling back to OCR")

    pages_filtered = False
    if has_text:
        # «Смешанный» PDF: часть страниц имеет нормальный текстовый слой (описание),
        # а ключевые таблицы характеристик нарисованы как векторная графика/картинка
        # без ToUnicode — их get_text() почти пуст, и LLM не видит значений. Для таких
        # «бедных» страниц прогоняем OCR и подменяем их текст распознанным.
        all_pages_text, augmented = _augment_sparse_pages_with_ocr(
            document, all_pages_text
        )
        if augmented:
            ocr_applied = True
            logger.info(
                "PDF mixed-content: OCR-augmented %d sparse page(s) with image/vector tables",
                augmented,
            )
        relevant_pages, pages_filtered = _filter_relevant_pages(
            all_pages_text, target_names
        )
        if pages_filtered:
            logger.info(
                "PDF page filter: %d/%d pages selected (total %d chars → %d chars)",
                len(relevant_pages), page_count,
                sum(len(t) for t in all_pages_text),
                sum(len(t) for _, t in relevant_pages),
            )
        for page_num, text in relevant_pages:
            lines.append(f"[PAGE {page_num}]")
            lines.append(text)
    else:
        ocr_available = False
        try:
            import pytesseract
            pytesseract.get_tesseract_version()
            ocr_available = True
        except Exception:
            pass

        if ocr_available:
            from PIL import Image
            import io
            DPI = 200
            try:
                import pytesseract
                for i in range(page_count):
                    page = document.load_page(i)
                    mat = fitz.Matrix(DPI / 72, DPI / 72)
                    pix = page.get_pixmap(matrix=mat, alpha=False)
                    img = Image.open(io.BytesIO(pix.tobytes("png")))
                    img = _auto_orient_for_ocr(img)
                    try:
                        text = pytesseract.image_to_string(
                            img, lang="rus+eng", timeout=OCR_PAGE_TIMEOUT_SECONDS
                        ).strip()
                    except RuntimeError as exc:
                        logger.warning("OCR timeout/error on page %d: %s", i + 1, exc)
                        text = ""
                    if text:
                        lines.append(f"[PAGE {i + 1}]")
                        lines.append(text)
                ocr_applied = True
            except Exception as exc:
                logger.warning("PDF OCR failed: %s", exc)

    document.close()
    raw_text = "\n".join(lines).strip()
    clean_text = raw_text.encode("utf-8", errors="replace").decode("utf-8")
    return {
        "text": clean_text,
        "page_count": page_count,
        "ocr_applied": ocr_applied,
        "pages_filtered": pages_filtered,
    }


def _guess_filename_from_url(file_url: str, content_type: Optional[str]) -> str:
    parsed_url = urlparse(file_url)
    raw_name = unquote(Path(parsed_url.path).name)
    if raw_name:
        return raw_name

    extension = mimetypes.guess_extension((content_type or "").split(";", 1)[0].strip())
    extension = extension or ".pdf"
    return f"document{extension}"


def _normalized_content_type(filename: str, content_type: Optional[str]) -> Optional[str]:
    normalized = (content_type or "").split(";", 1)[0].strip().lower()
    if normalized and normalized not in GENERIC_MIME_TYPES:
        return normalized
    guessed, _ = mimetypes.guess_type(filename)
    return guessed or normalized or None


def _message_text_from_content(content: Any) -> str:
    if isinstance(content, str):
        return content
    if isinstance(content, list):
        text_parts: list[str] = []
        for item in content:
            if isinstance(item, dict) and item.get("type") == "text":
                text = item.get("text")
                if isinstance(text, str):
                    text_parts.append(text)
        return "\n".join(part for part in text_parts if part)
    return ""


def _strip_json_fences(text: str) -> str:
    candidate = text.strip()
    if candidate.startswith("```"):
        lines = candidate.splitlines()
        if lines and lines[0].startswith("```"):
            lines = lines[1:]
        if lines and lines[-1].strip() == "```":
            lines = lines[:-1]
        candidate = "\n".join(lines).strip()
    return candidate


def _extract_message_json_text(data: dict[str, Any], *, provider_name: str) -> str:
    choices = data.get("choices")
    if not isinstance(choices, list) or not choices:
        raise RuntimeError(f"{provider_name} response has no choices")

    message = choices[0].get("message", {})
    content = _message_text_from_content(message.get("content"))
    content = _strip_json_fences(content)
    if not content:
        raise RuntimeError(f"{provider_name} returned empty content")
    return content


def _json_error_snippet(raw_content: str, position: int, radius: int = 700) -> str:
    start = max(0, position - radius)
    end = min(len(raw_content), position + radius)
    prefix = "... " if start > 0 else ""
    suffix = " ..." if end < len(raw_content) else ""
    return prefix + raw_content[start:end] + suffix


def _provider_response_metadata(data: dict[str, Any]) -> dict[str, Any]:
    metadata: dict[str, Any] = {}
    for key in ("id", "model", "provider"):
        value = data.get(key)
        if value is not None:
            metadata[key] = value
    choices = data.get("choices")
    if isinstance(choices, list) and choices and isinstance(choices[0], dict):
        choice = choices[0]
        for key in ("finish_reason", "native_finish_reason"):
            value = choice.get(key)
            if value is not None:
                metadata[key] = value
    usage = data.get("usage")
    if isinstance(usage, dict):
        metadata["usage"] = usage
    return metadata


def _json_decode_diagnostic(
    provider_name: str,
    raw_content: str,
    exc: json.JSONDecodeError,
    response_data: dict[str, Any],
) -> str:
    snippet = _json_error_snippet(raw_content, exc.pos)
    metadata = _provider_response_metadata(response_data)
    return (
        f"{provider_name} returned invalid JSON: {exc.msg} "
        f"at line={exc.lineno} column={exc.colno} char={exc.pos}; "
        f"response_length={len(raw_content)}; "
        f"response_metadata={json.dumps(metadata, ensure_ascii=True)}; "
        f"snippet_around_error={json.dumps(snippet, ensure_ascii=True)}"
    )


def _extract_error_text_from_response(response: httpx.Response, provider_name: str) -> str:
    try:
        payload = response.json()
    except ValueError:
        text = (response.text or "").strip()
        return text[:800] if text else f"{provider_name} returned HTTP {response.status_code}"

    if isinstance(payload, dict):
        error = payload.get("error")
        if isinstance(error, dict):
            message = error.get("message") or error.get("metadata")
            if isinstance(message, str) and message.strip():
                return message.strip()[:800]
            if message is not None:
                return str(message)[:800]
        detail = payload.get("detail")
        if isinstance(detail, str) and detail.strip():
            return detail.strip()[:800]
        if detail is not None:
            return str(detail)[:800]

    text = json.dumps(payload, ensure_ascii=False)[:800]
    return text or f"{provider_name} returned HTTP {response.status_code}"


def _raise_provider_http_error(response: httpx.Response, provider_name: str) -> None:
    detail = _extract_error_text_from_response(response, provider_name)
    raise RuntimeError(
        f"{provider_name} HTTP {response.status_code}: {detail}"
    )


def _json_schema_response_format(name: str, schema: dict[str, Any]) -> dict[str, Any]:
    if LLM_IS_YANDEX:
        return {"type": "json_object"}
    # strict=False: схема передаётся модели как ОРИЕНТИР, без жёсткой валидации
    # структуры. Это критично для gpt-4.1 (и новее): в strict-режиме OpenAI требует
    # 'additionalProperties': false и полный 'required' в КАЖДОМ объекте схемы,
    # а наша схема с references/bbox/anyOf этим требованиям не удовлетворяет —
    # strict-запрос отклонялся с 400, падал в fallback без схемы, и модель
    # возвращала голый массив (non-object JSON → весь анализ failed).
    # Со strict=False gpt-4.1 корректно возвращает {products: [...]}.
    return {
        "type": "json_schema",
        "json_schema": {
            "name": name,
            "strict": False,
            "schema": schema,
        },
    }


def _build_openrouter_provider_preferences(*, require_parameters: bool) -> dict[str, Any]:
    if LLM_IS_YANDEX:
        return {}
    provider: dict[str, Any] = {}
    if OPENROUTER_PROVIDER_ORDER:
        provider["order"] = OPENROUTER_PROVIDER_ORDER
    if OPENROUTER_PROVIDER_IGNORE:
        provider["ignore"] = OPENROUTER_PROVIDER_IGNORE
    if require_parameters:
        provider["require_parameters"] = True
    return provider


async def _download_file(file_url: str) -> DownloadedFile:
    temp_path = ""
    try:
        async with httpx.AsyncClient(
            timeout=FILE_DOWNLOAD_TIMEOUT_SECONDS,
            follow_redirects=True,
            trust_env=False,
        ) as client:
            async with client.stream("GET", file_url) as response:
                response.raise_for_status()
                content_type = response.headers.get("content-type")
                filename = _guess_filename_from_url(file_url, content_type)
                suffix = Path(filename).suffix or (
                    mimetypes.guess_extension((content_type or "").split(";", 1)[0].strip())
                    or ".pdf"
                )
                with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as handle:
                    temp_path = handle.name
                    async for chunk in response.aiter_bytes():
                        handle.write(chunk)
    except httpx.HTTPError as exc:
        raise RuntimeError(f"Failed to download file: {exc}") from exc
    except Exception:
        if temp_path:
            try:
                os.remove(temp_path)
            except OSError:
                logger.warning("Failed to remove temp file %s", temp_path)
        raise

    return DownloadedFile(
        filename=filename,
        content_type=_normalized_content_type(filename, content_type),
        local_path=temp_path,
    )


def _normalize_match_text(value: str | None) -> str:
    if not value:
        return ""
    return re.sub(r"\s+", " ", re.sub(r"[^\w\s./%-]+", " ", value.lower().replace("ё", "е"))).strip()


def _tokenize_match_text(value: str | None) -> list[str]:
    normalized = _normalize_match_text(value)
    if not normalized:
        return []
    return [token for token in normalized.split(" ") if len(token) >= 2]


def _matched_text_consistent_with_quote(
    matched_text: str | None, quote_text: str | None
) -> bool:
    """True, если matched_text согласуется с quote_text (один — подстрока другого).

    Зеркалит проверку фронтенда (isFallbackAnchor): если геометрия нашла текст,
    не совпадающий с цитатой, фронт прячет такую подсветку как «обманчивую».
    Возвращаем True только когда matched_text реально соответствует цитате —
    тогда его безопасно выставлять. Для таблиц (нашли строку по названию, а
    quote — это значение) вернём False, и matched_text не сохранится, чтобы
    корректный bbox не был отброшен фронтендом."""
    if not matched_text or not quote_text:
        return False
    norm_matched = _normalize_match_text(matched_text)
    norm_quote = _normalize_match_text(quote_text)
    if not norm_matched or not norm_quote:
        return False
    if norm_matched == norm_quote:
        return True
    return norm_quote in norm_matched or norm_matched in norm_quote


def _safe_page_number(value: Any) -> int | None:
    if value is None or value == "":
        return None
    if isinstance(value, int):
        return value
    if isinstance(value, float) and value.is_integer():
        return int(value)
    if isinstance(value, str) and value.strip().isdigit():
        return int(value.strip())
    return None


def _normalize_references_in_place(node: Any) -> None:
    """Нормализует поле references везде в дереве:
    если модель вернула dict вместо list — оборачивает в list.
    Некоторые модели (gpt-4o-mini) возвращают один объект вместо массива.
    """
    if isinstance(node, dict):
        references = node.get("references")
        if isinstance(references, dict):
            node["references"] = [references]
        for key, value in node.items():
            if key != "references":
                _normalize_references_in_place(value)
    elif isinstance(node, list):
        for value in node:
            _normalize_references_in_place(value)


def _reference_iter(node: Any):
    if isinstance(node, dict):
        references = node.get("references")
        if isinstance(references, list):
            yield node, references
        for key, value in node.items():
            if key == "references":
                continue
            yield from _reference_iter(value)
    elif isinstance(node, list):
        for value in node:
            yield from _reference_iter(value)


def _collect_generic_anchor_texts(root: Any) -> set[str]:
    """Находит «общие» anchor/locator-тексты, которые нельзя использовать как
    привязку к месту в PDF.

    Типичная проблема: LLM ставит одинаковый anchor_text (название изделия,
    например «Гидрант пожарный») для ВСЕХ характеристик. Геометрия тогда находит
    этот заголовок один раз на стр. 1 и сажает туда все характеристики.
    Считаем такими «общими» якорями: (1) названия изделий (product_name) и
    (2) anchor/locator-тексты, повторяющиеся в >=2 референсах."""
    generic: set[str] = set()
    counts: dict[str, int] = {}

    def walk(node: Any) -> None:
        if isinstance(node, dict):
            for key in ("product_name", "product_model"):
                value = node.get(key)
                if isinstance(value, str):
                    norm = _normalize_match_text(value)
                    if norm:
                        generic.add(norm)
            references = node.get("references")
            if isinstance(references, list):
                for reference in references:
                    if not isinstance(reference, dict):
                        continue
                    for key in ("anchor_text", "locator_text"):
                        value = reference.get(key)
                        if isinstance(value, str):
                            norm = _normalize_match_text(value)
                            if norm:
                                counts[norm] = counts.get(norm, 0) + 1
            for key, value in node.items():
                if key != "references":
                    walk(value)
        elif isinstance(node, list):
            for value in node:
                walk(value)

    walk(root)
    generic.update(text for text, count in counts.items() if count >= 2)
    return generic


def _reference_iter_with_ancestors(node: Any, ancestors: tuple[dict[str, Any], ...] = ()):
    if isinstance(node, dict):
        references = node.get("references")
        if isinstance(references, list):
            yield node, references, ancestors
        next_ancestors = (*ancestors, node)
        for key, value in node.items():
            if key == "references":
                continue
            yield from _reference_iter_with_ancestors(value, next_ancestors)
    elif isinstance(node, list):
        for value in node:
            yield from _reference_iter_with_ancestors(value, ancestors)


def _collect_ancestor_context(ancestors: tuple[dict[str, Any], ...]) -> list[str]:
    context_values: list[str] = []
    for node in reversed(ancestors):
        product_name = node.get("product_name")
        product_model = node.get("product_model")
        if isinstance(product_name, str) and product_name.strip():
            if isinstance(product_model, str) and product_model.strip():
                context_values.append(f"{product_name.strip()} {product_model.strip()}")
            context_values.append(product_name.strip())
        if isinstance(product_model, str) and product_model.strip():
            context_values.append(product_model.strip())
        for key in ("title", "label", "name"):
            value = node.get(key)
            if isinstance(value, str) and value.strip():
                context_values.append(value.strip())

    deduped: list[str] = []
    seen: set[str] = set()
    for value in context_values:
        normalized = _normalize_match_text(value)
        if not normalized or normalized in seen:
            continue
        seen.add(normalized)
        deduped.append(value)
    return deduped[:8]


def _build_synthetic_reference(label_context: list[str], value_context: list[str]) -> dict[str, Any] | None:
    primary_value = next((v for v in value_context if v), "")
    primary_label = " ".join(label_context[:2]).strip()
    if not primary_value and not primary_label:
        return None
    value_is_ambiguous = primary_value and _is_ambiguous_short_candidate(
        primary_value, _normalize_match_text(primary_value)
    )
    if value_is_ambiguous and primary_label:
        quote_text = f"{primary_label} {primary_value}"
        anchor_text = primary_label
    else:
        quote_text = primary_value or primary_label
        anchor_text = quote_text
    return {
        "quote_text": quote_text,
        "anchor_text": anchor_text,
        "locator_text": anchor_text,
        "synthetic_reference": True,
    }


def _inject_synthetic_references(node: Any) -> int:
    injected = 0
    if isinstance(node, dict):
        label_context = _collect_label_context(node)
        value_context = _collect_value_context(node)
        has_reference_key = "references" in node
        references = node.get("references")
        if (
            has_reference_key
            and isinstance(references, list)
            and not references
            and (label_context or value_context)
        ):
            ref = _build_synthetic_reference(label_context, value_context)
            if ref:
                node["references"] = [ref]
                injected += 1
        elif (
            not has_reference_key
            and label_context
            and value_context
            and any(key in node for key in ("name", "label", "title", "characteristic"))
            and any(key in node for key in ("value", "text", "answer", "content"))
        ):
            ref = _build_synthetic_reference(label_context, value_context)
            if ref:
                node["references"] = [ref]
                injected += 1
        for key, value in list(node.items()):
            if key == "references":
                continue
            injected += _inject_synthetic_references(value)
    elif isinstance(node, list):
        for value in node:
            injected += _inject_synthetic_references(value)
    return injected


def _collect_label_context(node: dict[str, Any]) -> list[str]:
    context_values: list[str] = []
    for key in (
        "name",
        "label",
        "title",
        "characteristic",
        "field_name",
        "product_name",
        "section",
        "key",
    ):
        value = node.get(key)
        if isinstance(value, str) and value.strip():
            context_values.append(value.strip())
    return context_values


def _collect_value_context(node: dict[str, Any]) -> list[str]:
    context_values: list[str] = []
    for key in ("value", "text", "answer", "content", "quote", "summary"):
        value = node.get(key)
        if isinstance(value, str) and value.strip():
            context_values.append(value.strip())
    return context_values


def _looks_like_markdown_table_row(text: str) -> bool:
    """True, если строка похожа на ячейки Markdown-таблицы (есть разделители «|»).
    LLM теперь получает таблицы как Markdown (см. _page_text_with_tables), поэтому
    quote_text характеристики из таблицы — это строка вида 'модель | v1 | v2 | ...'.
    В реальном тексте PDF символов «|» нет, и точный поиск такой цитаты проваливался
    → геометрия привязывалась к заголовку. Поэтому такие цитаты разбираем на ячейки."""
    return text.count("|") >= 2


def _markdown_row_cell_candidates(
    text: str, page_hint: int | None
) -> list[dict[str, Any]]:
    """Из Markdown-строки таблицы делает кандидаты для геометрии: код модели (первая
    значимая ячейка), отдельные значения-ячейки и очищенный текст без «|».
    Markdown-разделители «| --- |» игнорируются."""
    cells = [c.strip() for c in text.split("|")]
    cells = [c for c in cells if c and not re.fullmatch(r"[-:\s]+", c)]
    if not cells:
        return []
    out: list[dict[str, Any]] = []
    # Первая ячейка — обычно код модели / название строки: ценный якорь.
    out.append({
        "text": cells[0], "kind": "anchor_text", "weight": 0.9, "page_hint": page_hint,
    })
    # Остальные ячейки — значения; короткие числа ambiguous, но в паре с моделью
    # дают точную привязку строки.
    for cell in cells[1:]:
        if len(cell) >= 2:
            out.append({
                "text": cell, "kind": "value", "weight": 0.75, "page_hint": page_hint,
            })
    # Очищенная строка целиком (числа через пробел) — иногда совпадает с PDF-строкой.
    cleaned = " ".join(cells)
    if cleaned and cleaned != text:
        out.append({
            "text": cleaned, "kind": "text", "weight": 0.7, "page_hint": page_hint,
        })
    return out


def _collect_reference_candidates(
    reference: Any,
    *,
    label_context: list[str],
    value_context: list[str],
    generic_anchors: set[str] | None = None,
) -> list[dict[str, Any]]:
    raw_candidates: list[dict[str, Any]] = []
    generic_anchors = generic_anchors or set()
    page_hint = None
    if isinstance(reference, dict):
        page_hint = _safe_page_number(reference.get("page"))
        if page_hint is None:
            page_hint = _safe_page_number(reference.get("page_number"))
        for key in ("quote_text", "anchor_text", "locator_text", "text"):
            value = reference.get(key)
            if isinstance(value, str) and value.strip():
                # Пропускаем «общие» якоря (название изделия и т.п.): по ним нельзя
                # позиционировать конкретную характеристику — иначе все они сядут
                # в одно место (заголовок на стр. 1).
                if key in {"anchor_text", "locator_text"} and _normalize_match_text(value) in generic_anchors:
                    continue
                # Markdown-строка таблицы: «|» в реальном PDF нет — разбираем на ячейки
                # (код модели + значения), иначе точный поиск цитаты провалится и
                # геометрия привяжется к заголовку таблицы.
                if key == "quote_text" and _looks_like_markdown_table_row(value):
                    raw_candidates.extend(
                        _markdown_row_cell_candidates(value, page_hint)
                    )
                    continue
                raw_candidates.append(
                    {
                        "text": value.strip(),
                        "kind": key,
                        "weight": {
                            "quote_text": 1.0,
                            "anchor_text": 0.88,
                            "locator_text": 0.83,
                            "text": 0.8,
                        }.get(key, 0.75),
                        "page_hint": page_hint,
                    }
                )
    elif isinstance(reference, str) and reference.strip():
        raw_candidates.append(
            {
                "text": reference.strip(),
                "kind": "raw_reference",
                "weight": 0.92,
                "page_hint": page_hint,
            }
        )

    primary_label = " ".join(label_context[:2]).strip()
    primary_value = next((value for value in value_context if value), "")
    value_is_ambiguous = primary_value and _is_ambiguous_short_candidate(
        primary_value, _normalize_match_text(primary_value)
    )
    if primary_label and primary_value:
        raw_candidates.append(
            {
                "text": f"{primary_label} {primary_value}",
                "kind": "label_plus_value",
                "weight": 0.95 if value_is_ambiguous else 0.82,
                "page_hint": page_hint,
            }
        )
    if primary_value:
        raw_candidates.append(
            {
                "text": primary_value,
                "kind": "value",
                "weight": 0.5 if value_is_ambiguous else 0.8,
                "page_hint": page_hint,
            }
        )
    if primary_label:
        raw_candidates.append(
            {
                "text": primary_label,
                "kind": "label",
                "weight": 0.45,
                "page_hint": page_hint,
            }
        )

    deduped: list[dict[str, Any]] = []
    seen: set[str] = set()
    for candidate in raw_candidates:
        normalized = _normalize_match_text(candidate["text"])
        if not normalized or normalized in seen:
            continue
        seen.add(normalized)
        deduped.append({**candidate, "normalized": normalized})
    return deduped[:8]


def _page_index_rect(page_index: "PdfPageIndex") -> Any:
    """Система координат для нормализации bbox страницы: для OCR-индекса с
    авто-поворотом — прямоугольник выпрямленного изображения, иначе — page.rect."""
    if getattr(page_index, "page_rect", None) is not None:
        return page_index.page_rect
    return page_index.page.rect


def _rect_to_bbox(rect: Any, page_rect: Any = None) -> dict[str, float]:
    result: dict[str, float] = {
        "x": round(float(rect.x0), 3),
        "y": round(float(rect.y0), 3),
        "width": round(float(rect.x1 - rect.x0), 3),
        "height": round(float(rect.y1 - rect.y0), 3),
        "x0": round(float(rect.x0), 3),
        "y0": round(float(rect.y0), 3),
        "x1": round(float(rect.x1), 3),
        "y1": round(float(rect.y1), 3),
        "left": round(float(rect.x0), 3),
        "top": round(float(rect.y0), 3),
        "right": round(float(rect.x1), 3),
        "bottom": round(float(rect.y1), 3),
    }
    if page_rect is not None:
        pw = float(page_rect.width)
        ph = float(page_rect.height)
        if pw > 0 and ph > 0:
            result["norm_x0"] = round(min(1.0, max(0.0, float(rect.x0) / pw)), 6)
            result["norm_y0"] = round(min(1.0, max(0.0, float(rect.y0) / ph)), 6)
            result["norm_x1"] = round(min(1.0, max(0.0, float(rect.x1) / pw)), 6)
            result["norm_y1"] = round(min(1.0, max(0.0, float(rect.y1) / ph)), 6)
    return result


def _union_rects(rects: list[Any]) -> Any | None:
    if not rects:
        return None
    current = fitz.Rect(rects[0])
    for rect in rects[1:]:
        current.include_rect(rect)
    return current


def _build_pdf_page_index(local_path: str) -> tuple[Any, list[PdfPageIndex]]:
    document = fitz.open(local_path)
    pages: list[PdfPageIndex] = []
    for page_index in range(document.page_count):
        page = document.load_page(page_index)
        raw_words = page.get_text("words", sort=True)
        words: list[PdfWord] = []
        for raw_word in raw_words:
            text = str(raw_word[4]).strip()
            normalized = _normalize_match_text(text)
            if not normalized:
                continue
            words.append(
                PdfWord(
                    text=text,
                    normalized=normalized,
                    rect=fitz.Rect(raw_word[0], raw_word[1], raw_word[2], raw_word[3]),
                )
            )
        pages.append(PdfPageIndex(page_number=page_index + 1, page=page, words=words))
    return document, pages




def _build_pdf_page_index_ocr(local_path: str) -> tuple[Any, list[PdfPageIndex]]:
    """
    For scanned PDFs: renders each page as an image, runs Tesseract OCR,
    and builds PdfPageIndex from the recognized words.
    Word bboxes are in PDF-point coordinates (at DPI=150 scale).
    """
    import pytesseract
    from PIL import Image
    import io

    # DPI=150 даёт самые стабильные заголовки столбцов моделей ('50-110' и т.п.),
    # что критично для привязки ячеек таблиц к нужной модели. Более высокий DPI
    # иногда лучше читает плотные строки значений, но дробит/теряет заголовки
    # моделей — а без якоря модели вся табличная привязка рушится.
    DPI = 150
    document = fitz.open(local_path)
    pages: list[PdfPageIndex] = []

    for page_index in range(document.page_count):
        page = document.load_page(page_index)
        mat = fitz.Matrix(DPI / 72, DPI / 72)
        pix = page.get_pixmap(matrix=mat, alpha=False)
        img_bytes = pix.tobytes("png")
        img = Image.open(io.BytesIO(img_bytes))
        # Выпрямляем страницу (повёрнутые сканы) — иначе OCR читает текст боком/
        # вверх ногами и выдаёт мусор, а координаты не совпадут с вьювером.
        img = _auto_orient_for_ocr(img)
        # Прямоугольник системы координат = размер ВЫПРЯМЛЕННОГО изображения в
        # PDF-пунктах. Именно в этой системе фронт нормализует bbox.
        scale = 72.0 / DPI  # pixels -> PDF points
        page_rect = fitz.Rect(0, 0, img.width * scale, img.height * scale)

        try:
            ocr_data = pytesseract.image_to_data(
                img,
                lang="rus+eng",
                output_type=pytesseract.Output.DICT,
                timeout=OCR_PAGE_TIMEOUT_SECONDS,
            )
        except Exception as exc:
            logger.warning("OCR failed for page %d: %s", page_index + 1, exc)
            pages.append(
                PdfPageIndex(page_number=page_index + 1, page=page, words=[], page_rect=page_rect)
            )
            continue

        words: list[PdfWord] = []
        n = len(ocr_data["text"])
        for i in range(n):
            text = str(ocr_data["text"][i]).strip()
            conf = int(ocr_data["conf"][i])
            if not text or conf < 30:
                continue
            normalized = _normalize_match_text(text)
            if not normalized:
                continue
            x = ocr_data["left"][i] * scale
            y = ocr_data["top"][i] * scale
            w = ocr_data["width"][i] * scale
            h = ocr_data["height"][i] * scale
            rect = fitz.Rect(x, y, x + w, y + h)
            words.append(PdfWord(text=text, normalized=normalized, rect=rect))

        pages.append(
            PdfPageIndex(page_number=page_index + 1, page=page, words=words, page_rect=page_rect)
        )

    return document, pages

def _search_exact_candidate_rects(page: Any, candidate: str) -> list[Any]:
    snippet = candidate.strip()
    if len(snippet) < 3:
        return []
    if len(snippet) > 220:
        snippet = snippet[:220].rsplit(" ", 1)[0].strip() or snippet[:220]
    try:
        rects = page.search_for(snippet)
    except Exception:
        rects = []
    if rects:
        return sorted(rects, key=lambda rect: (float(rect.y0), float(rect.x0)))

    # Fallback: если точный поиск не нашёл (из-за множественных пробелов в таблицах),
    # нормализуем пробелы и ищем ещё раз
    normalized = re.sub(r"\s+", " ", snippet).strip()
    if normalized != snippet and len(normalized) >= 3:
        try:
            rects = page.search_for(normalized)
        except Exception:
            rects = []
        if rects:
            return sorted(rects, key=lambda rect: (float(rect.y0), float(rect.x0)))

    # Fallback 2: ищем только первую значимую часть (до первого большого пробела)
    parts = re.split(r"\s{2,}", snippet)
    if len(parts) > 1 and len(parts[0].strip()) >= 4:
        try:
            rects = page.search_for(parts[0].strip())
        except Exception:
            rects = []
        if rects:
            return sorted(rects, key=lambda rect: (float(rect.y0), float(rect.x0)))

    return []


def _rect_center(rect: Any) -> tuple[float, float]:
    return ((float(rect.x0) + float(rect.x1)) / 2, (float(rect.y0) + float(rect.y1)) / 2)


def _select_best_exact_rect(
    rects: list[Any],
    page_index: PdfPageIndex,
    context_terms: list[str],
) -> tuple[Any | None, float]:
    if not rects:
        return None, 0.0
    if len(rects) == 1:
        return rects[0], 0.0

    context_rects: list[Any] = []
    for term in context_terms:
        context_rects.extend(_search_exact_candidate_rects(page_index.page, term))
        if context_rects:
            break

    if not context_rects:
        return rects[0], 0.0

    def distance_to_context(rect: Any) -> float:
        rect_x, rect_y = _rect_center(rect)
        best = float("inf")
        for context_rect in context_rects:
            context_x, context_y = _rect_center(context_rect)
            vertical_distance = abs(rect_y - context_y)
            horizontal_distance = abs(rect_x - context_x) * 0.15
            best = min(best, vertical_distance + horizontal_distance)
        return best

    best_rect = min(rects, key=distance_to_context)
    context_bonus = max(0.0, 80.0 - distance_to_context(best_rect))
    return best_rect, context_bonus


def _bbox_is_compact(word_rects: list[Any], page_index: PdfPageIndex) -> bool:
    """Rejects a union bbox that spans too large a portion of the page.
    Typical table row is ~2-5% of page height and rarely the full width; a union
    spanning >12% height OR >85% width means matched tokens are scattered across
    multiple rows/columns → wrong match (e.g. value in one column, label in another)."""
    if not word_rects:
        return True
    page_rect = _page_index_rect(page_index)
    page_height = float(page_rect.height) if page_rect else 842.0  # A4 default
    page_width = float(page_rect.width) if page_rect else 595.0
    if page_height <= 0 or page_width <= 0:
        return True
    y_min = min(float(r.y0) for r in word_rects)
    y_max = max(float(r.y1) for r in word_rects)
    x_min = min(float(r.x0) for r in word_rects)
    x_max = max(float(r.x1) for r in word_rects)
    if (y_max - y_min) / page_height > 0.12:
        return False
    # Wide spans are only suspicious for multi-token matches: a single matched
    # phrase is naturally narrow; a >85%-width union means tokens jumped columns.
    if len(word_rects) >= 3 and (x_max - x_min) / page_width > 0.85:
        return False
    return True


def _search_token_candidate(page_index: PdfPageIndex, candidate: str) -> tuple[Any | None, float]:
    tokens = _tokenize_match_text(candidate)
    if len(tokens) < 2 or not page_index.words:
        return None, 0.0

    best_indices: list[int] = []
    best_coverage = 0.0
    max_window = min(len(page_index.words), max(len(tokens) + 12, 18))

    for start in range(len(page_index.words)):
        seen_indices: list[int] = []
        token_cursor = 0
        for index in range(start, min(len(page_index.words), start + max_window)):
            if token_cursor >= len(tokens):
                break
            word = page_index.words[index].normalized
            if word == tokens[token_cursor]:
                seen_indices.append(index)
                token_cursor += 1
                continue
            if token_cursor < len(tokens) and tokens[token_cursor] in word:
                seen_indices.append(index)
                token_cursor += 1
        matched_count = len(seen_indices)
        coverage = matched_count / len(tokens)
        if matched_count > len(best_indices) or (
            matched_count == len(best_indices) and coverage > best_coverage
        ):
            best_indices = seen_indices
            best_coverage = coverage

    matched_count = len(best_indices)
    if matched_count < min(3, len(tokens)) and best_coverage < 0.75:
        return None, 0.0
    if best_coverage < 0.55:
        return None, 0.0

    word_rects = [page_index.words[index].rect for index in best_indices]
    if not _bbox_is_compact(word_rects, page_index):
        return None, 0.0

    rect = _union_rects(word_rects)
    return rect, best_coverage


def _search_fuzzy_token_candidate(
    page_index: PdfPageIndex, candidate: str
) -> tuple[Any | None, float]:
    """OCR-устойчивый поиск с привязкой к ОДНОМУ компактному месту на странице.

    Берём только различительные токены цитаты (длиной >= 4: имена характеристик,
    а не служебные «не»/«до» и не голые числа, которые встречаются по всей странице),
    и ищем тесное непрерывное окно слов, где встречается большинство из них.
    Так мы избегаем ложных совпадений «по разбросанным общим словам», которые
    давали одинаковый счёт на разных страницах. Возвращает прямоугольник найденного
    места и долю совпавших различительных токенов."""
    distinctive = [token for token in _tokenize_match_text(candidate) if len(token) >= 4]
    if len(distinctive) < 2 or not page_index.words:
        return None, 0.0

    token_set = set(distinctive)
    word_norms = [word.normalized for word in page_index.words]

    def token_hits(word: str) -> str | None:
        for token in token_set:
            if word == token or (len(word) >= 4 and (token in word or word in token)):
                return token
        return None

    # Узкое окно: совпадения должны идти кучно (одно место в документе),
    # а не быть рассыпаны по всей странице.
    window = len(distinctive) + 3
    best_indices: list[int] = []
    best_ratio = 0.0
    for start in range(len(page_index.words)):
        end = min(len(page_index.words), start + window)
        hit_indices: list[int] = []
        matched_tokens: set[str] = set()
        for index in range(start, end):
            token = token_hits(word_norms[index])
            if token is not None and token not in matched_tokens:
                matched_tokens.add(token)
                hit_indices.append(index)
        ratio = len(matched_tokens) / len(token_set)
        if ratio > best_ratio or (ratio == best_ratio and len(hit_indices) < len(best_indices)):
            best_ratio = ratio
            best_indices = hit_indices

    # Требуем уверенное совпадение: >=70% различительных токенов И минимум 2 из них.
    if best_ratio < 0.7 or len(best_indices) < 2:
        return None, 0.0

    word_rects = [page_index.words[index].rect for index in best_indices]
    if not _bbox_is_compact(word_rects, page_index):
        return None, 0.0

    rect = _union_rects(word_rects)
    return rect, best_ratio


def _token_overlap_ratio(a: str, b: str) -> float:
    """Fraction of tokens from a that are found in b."""
    tokens_a = set(_tokenize_match_text(a))
    tokens_b = set(_tokenize_match_text(b))
    if not tokens_a:
        return 0.0
    return len(tokens_a & tokens_b) / len(tokens_a)


_tables_cache: dict[int, Any] = {}


def _get_page_tables(page_index: PdfPageIndex) -> Any:
    page_id = id(page_index.page)
    cached = _tables_cache.get(page_id)
    if cached is not None:
        return cached
    try:
        tables = page_index.page.find_tables()
    except Exception:
        tables = None
    _tables_cache[page_id] = tables
    return tables


def _model_numeric_cores(text: str) -> list[str]:
    """Числовое ядро типоразмера: '5Кс — 5х4 (КС 50-110/4)' → ['50-110/4','50-110'].
    Паспорт пишет модель как '1Кс50-110' — полную строку из ТЗ search_for не найдёт,
    а ядро '50-110' совпадает точно."""
    norm = text.lower().replace("–", "-").replace("—", "-").replace("х", "x")
    cores: list[str] = []
    for m in re.findall(r"\d+(?:[-/x]\d+)+", norm):
        if m not in cores:
            cores.append(m)
        base = m.split("/", 1)[0]
        if "-" in base and base not in cores:
            cores.append(base)
    return cores


def _model_term_rects(page_index: PdfPageIndex, context_terms: list[str]) -> list[Any]:
    """Прямоугольники, где на странице встречается код модели из контекста —
    по полной строке И по числовому ядру типоразмера (устойчиво к разным префиксам)."""
    rects: list[Any] = []
    search_terms: list[str] = []
    for term in context_terms:
        norm_term = _normalize_match_text(term)
        if not norm_term or len(norm_term) < 4:
            continue
        search_terms.append(term)
        # Числовое ядро (50-110) — паспорт обычно пишет модель именно так.
        for core in _model_numeric_cores(term):
            if len(core) >= 4 and core not in search_terms:
                search_terms.append(core)
    for term in search_terms:
        rects.extend(_search_exact_candidate_rects(page_index.page, term))
    return rects


def _model_rects_from_words(page_index: PdfPageIndex, context_terms: list[str]) -> list[Any]:
    """Прямоугольники кода модели, найденные по СЛОВАМ страницы (page_index.words).

    В отличие от _model_term_rects (через page.search_for), работает и на OCR-страницах,
    где реального текстового слоя нет. Ищем числовое ядро типоразмера ('50-110') как
    подстроку в нормализованном тексте каждого слова — паспорт пишет заголовок столбца
    модели как '50-110', '1Кс50-110', '50-110-...' и т.п."""
    cores: list[str] = []
    for term in context_terms:
        for core in _model_numeric_cores(term):
            norm_core = _normalize_match_text(core)
            if len(norm_core) >= 4 and norm_core not in cores:
                cores.append(norm_core)
    if not cores:
        return []

    rects: list[Any] = []
    matched_y: list[float] = []
    for word in page_index.words:
        wn = word.normalized
        for core in cores:
            if core in wn:
                rects.append(word.rect)
                matched_y.append(float(word.rect.y0))
                break
    if rects:
        return rects

    # OCR мог раздробить заголовок '50-110' на соседние слова '50' и '110'
    # (или '50-' + '110'). Восстанавливаем: для каждого ядра вида 'A-B' ищем
    # слово, начинающееся на A, и рядом (та же строка, правее, близко по X)
    # слово, заканчивающееся на B. Прямоугольник — объединение двух слов.
    for core in cores:
        parts = core.split("-")
        if len(parts) != 2 or not (parts[0].isdigit() and parts[1].isdigit()):
            continue
        a, b = parts[0], parts[1]
        a_words = [w for w in page_index.words if w.normalized.rstrip("-") == a or w.normalized == a]
        b_words = [w for w in page_index.words if w.normalized == b or w.normalized.lstrip("-") == b]
        for aw in a_words:
            ay = (float(aw.rect.y0) + float(aw.rect.y1)) / 2
            ah = float(aw.rect.y1) - float(aw.rect.y0)
            for bw in b_words:
                by = (float(bw.rect.y0) + float(bw.rect.y1)) / 2
                same_row = abs(ay - by) <= max(ah, 6.0)
                right_after = 0 <= (float(bw.rect.x0) - float(aw.rect.x1)) <= max(ah * 2, 20.0)
                if same_row and right_after:
                    rects.append(fitz.Rect(
                        min(float(aw.rect.x0), float(bw.rect.x0)),
                        min(float(aw.rect.y0), float(bw.rect.y0)),
                        max(float(aw.rect.x1), float(bw.rect.x1)),
                        max(float(aw.rect.y1), float(bw.rect.y1)),
                    ))
    return rects


def _value_cells_in_table(page_index: PdfPageIndex, norm_value: str) -> list[Any]:
    """Все ячейки таблиц, текст которых равен искомому значению."""
    tables = _get_page_tables(page_index)
    if not tables:
        return []
    cells: list[Any] = []
    for table in tables.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell is None:
                    continue
                try:
                    cell_rect = fitz.Rect(cell) if not isinstance(cell, fitz.Rect) else cell
                except Exception:
                    continue
                cell_text = page_index.page.get_textbox(cell_rect).strip()
                if _normalize_match_text(cell_text) == norm_value:
                    cells.append(cell_rect)
    return cells


def _search_value_in_model_table(
    page_index: PdfPageIndex,
    value_text: str,
    context_terms: list[str],
    param_label: str | None = None,
) -> tuple[Any | None, float]:
    """Находит ячейку со значением `value_text` в модельной таблице.

    Поддерживает две раскладки:
      • обычная (строка = модель): значение на той же ГОРИЗОНТАЛИ, что и код модели;
      • транспонированная (столбец = модель, строка = параметр L/B/H): значение на
        пересечении СТОЛБЦА модели (по X) и СТРОКИ параметра (по Y).

    Anchor вида "Мощность, кВт 30" или "Типоразмер насоса 1Кс 50-110: L 1195" нельзя
    искать целиком — заголовок и значение лежат в разных ячейках. Ищем ровно ячейку
    значения, привязанную к строке/столбцу модели."""
    norm_value = _normalize_match_text(value_text)
    if not norm_value:
        return None, 0.0

    # Работаем НАПРЯМУЮ по словам страницы (page_index.words), а не через
    # page.search_for()/find_tables — на OCR-страницах (сканы) реального текстового
    # слоя нет, и оба не работают. Слова же есть всегда (из текстового слоя или OCR),
    # поэтому word-based поиск универсален.
    model_rects = _model_rects_from_words(page_index, context_terms)
    if not model_rects:
        return None, 0.0
    # Ячейки-кандидаты со значением — слова, чей нормализованный текст равен искомому.
    value_cells = [w.rect for w in page_index.words if w.normalized == norm_value]
    if not value_cells:
        return None, 0.0

    def _center(rect: Any) -> tuple[float, float]:
        return ((float(rect.x0) + float(rect.x1)) / 2, (float(rect.y0) + float(rect.y1)) / 2)

    # Высота строки модели — масштаб допусков (OCR-слова мельче ячеек таблиц).
    model_h = max(
        (float(mr.y1) - float(mr.y0) for mr in model_rects), default=12.0
    ) or 12.0
    y_tol = max(model_h * 0.8, 8.0)

    def _layout_same_row() -> Any | None:
        """Значение на ОДНОЙ СТРОКЕ с кодом модели (обычная раскладка)."""
        cell, dist = None, float("inf")
        for cell_rect in value_cells:
            _, cy = _center(cell_rect)
            for mr in model_rects:
                y0, y1 = float(mr.y0) - y_tol, float(mr.y1) + y_tol
                d = 0.0 if y0 <= cy <= y1 else min(abs(cy - y0), abs(cy - y1))
                if d < dist:
                    dist, cell = d, cell_rect
        return cell if (cell is not None and dist <= y_tol) else None

    def _layout_transposed() -> Any | None:
        """Транспонированная: пересечение СТОЛБЦА модели (X) и СТРОКИ параметра (Y).

        Выбираем ячейку значения, X которой ближе всего к X-столбцу модели, а Y —
        к Y-строке параметра. Это устраняет неоднозначность, когда одно и то же число
        (напр. '280') встречается в нескольких строках/столбцах."""
        if not param_label:
            return None
        norm_label = _normalize_match_text(param_label)
        # Для односимвольных меток (L/B/H) — точное совпадение (иначе 'l' ловит всё).
        # Для словесных ('масса') — точное ИЛИ префиксное (OCR склеивает 'Масса,кг').
        if len(norm_label) <= 2:
            label_words = [w for w in page_index.words if w.normalized == norm_label]
        else:
            label_words = [
                w for w in page_index.words
                if w.normalized == norm_label or w.normalized.startswith(norm_label)
            ]
        if not label_words:
            return None
        # Метка строки — самая левая (это подпись строки, а не значение в ячейке).
        leftmost_x = min(float(w.rect.x0) for w in label_words)
        param_rects = [w.rect for w in label_words if float(w.rect.x0) <= leftmost_x + 40]
        if not param_rects:
            return None
        model_xs = [(float(mr.x0) + float(mr.x1)) / 2 for mr in model_rects]
        x_tol = max(model_h * 4, 60.0)
        cell, best = None, float("inf")
        for cell_rect in value_cells:
            cx, cy = _center(cell_rect)
            x_dist = min(abs(cx - mx) for mx in model_xs)
            y_dist = min(
                0.0 if (float(pr.y0) - y_tol) <= cy <= (float(pr.y1) + y_tol)
                else min(abs(cy - float(pr.y0)), abs(cy - float(pr.y1)))
                for pr in param_rects
            )
            if y_dist <= y_tol and x_dist <= x_tol:
                score = y_dist * 3 + x_dist
                if score < best:
                    best, cell = score, cell_rect
        return cell

    def _layout_column_only() -> tuple[Any | None, float]:
        """Fallback: значение в СТОЛБЦЕ модели (ближайшее по X к заголовку модели),
        НИЖЕ заголовка. Применяется, когда метку строки (Масса/L/B/H) OCR не распознал —
        тогда нельзя выбрать точную строку, но можно хотя бы попасть в нужный столбец,
        что устраняет грубую ошибку (значение из чужой модели). Уверенность ниже."""
        model_xs = [(float(mr.x0) + float(mr.x1)) / 2 for mr in model_rects]
        header_bottom = max(float(mr.y1) for mr in model_rects)
        # Узкий допуск по X — иначе залезаем в соседнюю колонку модели. Колонка
        # обычно шире метки модели; берём ~ширину заголовка модели как полупорог.
        model_w = max((float(mr.x1) - float(mr.x0) for mr in model_rects), default=40.0)
        x_tol = max(model_w * 0.7, 25.0)
        cell, best = None, float("inf")
        for cell_rect in value_cells:
            cx, cy = _center(cell_rect)
            if cy < header_bottom:  # выше заголовка модели — не данные этой таблицы
                continue
            x_dist = min(abs(cx - mx) for mx in model_xs)
            if x_dist <= x_tol and x_dist < best:
                best, cell = x_dist, cell_rect
        return (cell, 0.5) if cell is not None else (None, 0.0)

    # Для габаритов и веса (есть метка строки) сначала пробуем транспонированную
    # раскладку — иначе значение случайно «прилипает» к заголовку модели по Y.
    if param_label:
        cell = _layout_transposed() or _layout_same_row()
    else:
        cell = _layout_same_row() or _layout_transposed()
    if cell is not None:
        return cell, 0.88
    # Метку строки не нашли (плохой OCR) — пробуем хотя бы попасть в столбец модели.
    if param_label:
        return _layout_column_only()
    return None, 0.0


_DIMENSION_LABELS = {
    "длина": ["l", "длина", "дл"],
    "ширина": ["b", "в", "ширина", "шир"],
    "высота": ["h", "н", "высота", "выс"],
}


def _dimension_param_label(anchor_text: str) -> str | None:
    """Из имени характеристики габарита возвращает короткую метку строки таблицы
    (L/B/H), по которой ищется строка параметра в транспонированной таблице.
    "Габаритные размеры: Длина" → 'L'."""
    norm = _normalize_match_text(anchor_text)
    for key, variants in _DIMENSION_LABELS.items():
        if key in norm:
            return variants[0].upper()
    return None


def _weight_param_label(anchor_text: str) -> str | None:
    """Для характеристик веса/массы насоса возвращает метку строки 'масса' —
    в таблицах габаритов масса идёт отдельной строкой 'Масса, кг', значения которой
    разнесены по столбцам моделей. Это позволяет искать ячейку массы на пересечении
    строки 'Масса' и столбца нужной модели, а не первое попавшееся число.

    Только для веса/массы НАСОСА — не для электродвигателя/плиты/агрегата, чьи
    значения в этой таблице отсутствуют."""
    norm = _normalize_match_text(anchor_text)
    if "электродвигател" in norm or "плит" in norm or "агрегат" in norm or "общий" in norm:
        return None
    if "вес насоса" in norm or "масса насоса" in norm or norm in {"вес насоса", "масса насоса"}:
        return "масса"
    # "Вес: Насоса" → нормализуется в "вес насоса"
    if ("вес" in norm or "масс" in norm) and "насос" in norm:
        return "масса"
    return None


def _search_table_row_candidate(
    page_index: PdfPageIndex,
    anchor_text: str,
    context_terms: list[str] | None = None,
    value_text: str | None = None,
) -> tuple[Any | None, float]:
    """
    Searches for the table row or text line containing anchor_text.
    Strategy:
      0. If value + model context given: find the value cell on the model's row.
      1. Use page.find_tables() — if anchor_text matches a cell, return bbox of the whole row.
      2. Otherwise find the text line whose words best overlap anchor tokens (Y-band grouping).
    """
    normalized_anchor = _normalize_match_text(anchor_text)
    if not normalized_anchor or len(normalized_anchor) < 2:
        return None, 0.0

    # --- Step 0: model-keyed table — find the value cell on the model's row/column ---
    if value_text and context_terms:
        norm_val = _normalize_match_text(value_text)
        if norm_val and _is_ambiguous_short_candidate(value_text, norm_val):
            param_label = _dimension_param_label(anchor_text)
            rect, conf = _search_value_in_model_table(
                page_index, value_text, context_terms, param_label=param_label
            )
            if rect is not None:
                return rect, conf

    # --- Step 1: Search in tables ---
    try:
        tables = _get_page_tables(page_index)
        for table in (tables.tables if tables else []):
            for row in table.rows:
                for cell in row.cells:
                    if cell is None:
                        continue
                    try:
                        cell_rect = fitz.Rect(cell) if not isinstance(cell, fitz.Rect) else cell
                    except Exception:
                        continue
                    cell_text = page_index.page.get_textbox(cell_rect).strip()
                    cell_normalized = _normalize_match_text(cell_text)
                    if not cell_normalized:
                        continue
                    if (
                        normalized_anchor in cell_normalized
                        or cell_normalized in normalized_anchor
                        or _token_overlap_ratio(normalized_anchor, cell_normalized) >= 0.6
                    ):
                        row_cells = [
                            fitz.Rect(c) if not isinstance(c, fitz.Rect) else c
                            for c in row.cells
                            if c is not None
                        ]
                        row_rect = _union_rects(row_cells)
                        if not row_rect or row_rect.is_empty:
                            continue
                        page_rect = _page_index_rect(page_index)
                        ph = float(page_rect.height) if page_rect else 842.0
                        pw = float(page_rect.width) if page_rect else 595.0
                        # Reject oversized rows (merged cells across visual rows)
                        if ph > 0 and float(row_rect.height) / ph > 0.12:
                            continue
                        overlap = _token_overlap_ratio(normalized_anchor, cell_normalized)
                        # For very wide rows (multi-column tables like "Размеры в мм"),
                        # the whole-row highlight is misleading — narrow to the matched
                        # cell so the user sees the actual value, not the full table width.
                        if pw > 0 and float(row_rect.width) / pw > 0.6:
                            return cell_rect, max(0.7, overlap)
                        return row_rect, max(0.7, overlap)
    except Exception as exc:
        logger.debug("find_tables failed on page %d: %s", page_index.page_number, exc)

    # --- Step 2: Free-text Y-band search ---
    if not page_index.words:
        return None, 0.0

    line_tolerance = 4.0  # words within this y0 delta are on the same line
    lines: list[list[PdfWord]] = []
    current_line: list[PdfWord] = []
    current_y = None
    for word in page_index.words:
        word_y = float(word.rect.y0)
        if current_y is None or abs(word_y - current_y) <= line_tolerance:
            current_line.append(word)
            if current_y is None:
                current_y = word_y
        else:
            if current_line:
                lines.append(current_line)
            current_line = [word]
            current_y = word_y
    if current_line:
        lines.append(current_line)

    anchor_tokens = set(_tokenize_match_text(anchor_text))
    if not anchor_tokens:
        return None, 0.0

    best_line = None
    best_overlap = 0.0
    for line in lines:
        line_tokens = set(w.normalized for w in line)
        overlap = len(anchor_tokens & line_tokens) / len(anchor_tokens)
        if overlap > best_overlap:
            best_overlap = overlap
            best_line = line

    if best_line and best_overlap >= 0.5:
        rect = _union_rects([w.rect for w in best_line])
        return rect, best_overlap

    return None, 0.0

def _is_ambiguous_short_candidate(text: str, normalized: str) -> bool:
    """Returns True if candidate is short or purely numeric/units."""
    if len(normalized) <= 6:
        return True
    if re.fullmatch(r"[\d.,\-+/%\s]+", text.strip()):
        return True
    return False


def _candidate_rank(candidate: dict[str, Any], has_descriptive_candidate: bool = False) -> float:
    base = {
        "quote_text": 1000.0,
        "label_plus_value": 950.0,
        "value": 900.0,
        "raw_reference": 800.0,
        "locator_text": 550.0,
        "text": 520.0,
        "anchor_text": 450.0,
        "label": 250.0,
    }.get(str(candidate.get("kind") or ""), 300.0)

    text = str(candidate.get("text") or "")
    normalized = str(candidate.get("normalized") or _normalize_match_text(text))

    # Ambiguous short/numeric candidates are unreliable — they match everywhere
    # in multi-page documents. Demote them heavily so that descriptive candidates
    # (label_plus_value, anchor_text) win.
    if candidate.get("kind") in {"quote_text", "value"}:
        if _is_ambiguous_short_candidate(text, normalized):
            return 150.0 if has_descriptive_candidate else 200.0
    return base


def _is_generic_table_anchor(candidate: dict[str, Any]) -> bool:
    kind = str(candidate.get("kind") or "")
    if kind not in {"anchor_text", "locator_text", "text"}:
        return False
    normalized = str(candidate.get("normalized") or _normalize_match_text(str(candidate.get("text") or "")))
    return bool(re.fullmatch(r"(таблица|table)\s*\d*", normalized, flags=re.IGNORECASE))


def _page_has_context(page_index: PdfPageIndex, context_terms: list[str]) -> bool:
    """Checks whether any of the context terms appear on the given page."""
    if not context_terms or not page_index.words:
        return False
    page_text_set = {w.normalized for w in page_index.words}
    for term in context_terms:
        term_tokens = _tokenize_match_text(term)
        if not term_tokens:
            continue
        if sum(1 for t in term_tokens if t in page_text_set) >= max(1, len(term_tokens) // 2):
            return True
    return False


def _find_reference_location(
    pages: list[PdfPageIndex],
    candidates: list[dict[str, Any]],
    context_terms: list[str] | None = None,
) -> dict[str, Any] | None:
    best_match: dict[str, Any] | None = None
    context_terms = context_terms or []
    has_specific_candidate = any(
        candidate.get("kind") in {"quote_text", "value", "label_plus_value", "raw_reference"}
        for candidate in candidates
    )
    has_descriptive_candidate = any(
        c.get("kind") in {"anchor_text", "locator_text", "label_plus_value"}
        and len(str(c.get("normalized") or _normalize_match_text(str(c.get("text") or "")))) > 6
        for c in candidates
    )

    def _sorted_pages(page_hint: int | None, nearby_only: bool = False) -> list[PdfPageIndex]:
        ordered = sorted(
            pages,
            key=lambda page_index: (
                0 if page_hint and page_index.page_number == page_hint else 1,
                abs(page_index.page_number - page_hint) if page_hint else 0,
                page_index.page_number,
            ),
        )
        if nearby_only and page_hint:
            return [p for p in ordered if abs(p.page_number - page_hint) <= 3]
        return ordered

    def _ambiguity_penalty(candidate: dict[str, Any], page_index: PdfPageIndex) -> float:
        """Heavy penalty for ambiguous short candidates that land on a page
        where none of the context terms (characteristic name, product name) appear.
        This prevents "280" from matching on page 58 when the characteristic
        "Масса насоса" is only on page 6."""
        text = str(candidate.get("text") or "")
        normalized = str(candidate.get("normalized") or _normalize_match_text(text))
        if not _is_ambiguous_short_candidate(text, normalized):
            return 0.0
        if not context_terms:
            return 0.0
        if _page_has_context(page_index, context_terms):
            return 0.0
        return -500.0

    any_has_page_hint = any(_safe_page_number(c.get("page_hint")) for c in candidates)
    large_doc = len(pages) > 10

    # --- Pass 0: модель-aware поиск ячейки значения для ГАБАРИТОВ ---
    # Габаритные quote вида "Типоразмер насоса 1Кс 50-110: L 1195" иначе ловятся
    # token-поиском за заголовок "Типоразмер насоса". Если есть модель и метка
    # габарита (L/B/H), приоритетно ищем ячейку значения в строке параметра ×
    # столбце модели — это даёт точную ячейку, а не заголовок.
    value_cand = next((c for c in candidates if c.get("kind") == "value"), None)
    anchor_cand = next(
        (c for c in candidates if c.get("kind") in {"anchor_text", "label", "label_plus_value"}),
        None,
    )
    if value_cand and context_terms:
        param_label = None
        for c in candidates:
            lbl = _dimension_param_label(str(c.get("text") or ""))
            if lbl:
                param_label = lbl
                break
        # Контекст характеристики (название) тоже может содержать метку — для веса
        # ("Вес: Насоса") она не лежит в candidates, поэтому проверяем context_terms.
        if param_label is None:
            for term in context_terms:
                lbl = _dimension_param_label(term) or _weight_param_label(term)
                if lbl:
                    param_label = lbl
                    break
        if param_label:
            page_hint = _safe_page_number(value_cand.get("page_hint"))
            for page_index in _sorted_pages(page_hint, nearby_only=bool(page_hint and large_doc)):
                rect, conf = _search_value_in_model_table(
                    page_index, value_cand["text"], context_terms, param_label=param_label
                )
                if rect is not None:
                    # Высокий score, чтобы token-поиск не перебил точную ячейку
                    # заголовком "Типоразмер насоса".
                    best_match = {
                        "page": page_index.page_number,
                        "bbox": _rect_to_bbox(rect, _page_index_rect(page_index)),
                        "score": 1300.0 + conf * 100,
                        "locator_strategy": "table_cell",
                        "matched_text": value_cand["text"],
                    }
                    break

    def _do_exact_pass(page_list_fn):
        nonlocal best_match
        for candidate in candidates:
            if has_specific_candidate and _is_generic_table_anchor(candidate):
                continue
            candidate_text = candidate["text"]
            candidate_weight = float(candidate.get("weight", 0.5))
            page_hint = _safe_page_number(candidate.get("page_hint"))
            for page_index in page_list_fn(page_hint):
                rects = _search_exact_candidate_rects(page_index.page, candidate_text)
                rect, context_bonus = _select_best_exact_rect(rects, page_index, context_terms)
                if not rect:
                    continue
                page_bonus = 18.0 if page_hint and page_index.page_number == page_hint else 0.0
                score = (
                    _candidate_rank(candidate, has_descriptive_candidate=has_descriptive_candidate)
                    + len(candidate.get("normalized", _normalize_match_text(candidate_text))) * candidate_weight
                    + page_bonus
                    + context_bonus
                    + _ambiguity_penalty(candidate, page_index)
                )
                if not best_match or score > best_match["score"]:
                    best_match = {
                        "page": page_index.page_number,
                        "bbox": _rect_to_bbox(rect, _page_index_rect(page_index)),
                        "score": score,
                        "locator_strategy": "pymupdf_exact",
                        "matched_text": candidate_text,
                    }

    def _do_token_pass(page_list_fn):
        nonlocal best_match
        for candidate in candidates:
            if has_specific_candidate and _is_generic_table_anchor(candidate):
                continue
            candidate_text = candidate["text"]
            candidate_weight = float(candidate.get("weight", 0.5))
            page_hint = _safe_page_number(candidate.get("page_hint"))
            for page_index in page_list_fn(page_hint):
                rect, coverage = _search_token_candidate(page_index, candidate_text)
                if not rect:
                    continue
                page_bonus = 12.0 if page_hint and page_index.page_number == page_hint else 0.0
                score = (
                    _candidate_rank(candidate, has_descriptive_candidate=has_descriptive_candidate)
                    + coverage * 100 * candidate_weight
                    + min(len(candidate_text), 120) / 10
                    + page_bonus
                    + _ambiguity_penalty(candidate, page_index)
                )
                if not best_match or score > best_match["score"]:
                    best_match = {
                        "page": page_index.page_number,
                        "bbox": _rect_to_bbox(rect, _page_index_rect(page_index)),
                        "score": score,
                        "locator_strategy": "pymupdf_tokens",
                        "matched_text": candidate_text,
                    }

    # --- Pass 1+2: exact + token search ---
    # For large docs with page hints: try nearby pages first, expand only if needed
    if large_doc and any_has_page_hint:
        _do_exact_pass(lambda hint: _sorted_pages(hint, nearby_only=True))
        _do_token_pass(lambda hint: _sorted_pages(hint, nearby_only=True))
        if not best_match or best_match["score"] <= 500:
            _do_exact_pass(lambda hint: _sorted_pages(hint))
            _do_token_pass(lambda hint: _sorted_pages(hint))
    else:
        _do_exact_pass(lambda hint: _sorted_pages(hint))
        _do_token_pass(lambda hint: _sorted_pages(hint))

    # --- Pass 3: table row / free-text line search ---
    # find_tables() is expensive on large PDFs — skip if passes 1-2 found a confident match
    if not best_match or best_match["score"] <= 500:
        anchor_candidates = [
            c for c in candidates
            if c.get("kind") in {"anchor_text", "locator_text", "label", "label_plus_value"}
            and len(c.get("normalized", "")) >= 3
        ]
        # Значение характеристики (для модельных таблиц: ищем ячейку значения на
        # строке нужной модели, а не заголовок в чужой колонке).
        value_candidate = next(
            (c["text"] for c in candidates if c.get("kind") == "value"), None
        )
        for candidate in anchor_candidates:
            if has_specific_candidate and _is_generic_table_anchor(candidate):
                continue
            candidate_text = candidate["text"]
            candidate_weight = float(candidate.get("weight", 0.5))
            page_hint = _safe_page_number(candidate.get("page_hint"))
            for page_index in _sorted_pages(page_hint):
                rect, overlap = _search_table_row_candidate(
                    page_index,
                    candidate_text,
                    context_terms=context_terms,
                    value_text=value_candidate,
                )
                if not rect:
                    continue
                page_bonus = 10.0 if page_hint and page_index.page_number == page_hint else 0.0
                score = (
                    _candidate_rank(candidate, has_descriptive_candidate=False)
                    + overlap * 80 * candidate_weight
                    + min(len(candidate_text), 120) / 10
                    + page_bonus
                )
                if not best_match or score > best_match["score"]:
                    best_match = {
                        "page": page_index.page_number,
                        "bbox": _rect_to_bbox(rect, _page_index_rect(page_index)),
                        "score": score,
                        "locator_strategy": "table_row",
                        "matched_text": candidate_text,
                    }

    # --- Pass 4: OCR-robust fuzzy search ---
    if best_match is None:
        fuzzy_candidates = [
            c
            for c in candidates
            if c.get("kind") in {"quote_text", "value", "label_plus_value", "raw_reference"}
            and not _is_generic_table_anchor(c)
        ]
        for candidate in fuzzy_candidates:
            candidate_text = candidate["text"]
            page_ratios: list[tuple[float, Any, Any]] = []
            for page_index in pages:
                rect, ratio = _search_fuzzy_token_candidate(page_index, candidate_text)
                if rect:
                    page_ratios.append((ratio, rect, page_index))
            if not page_ratios:
                continue
            page_ratios.sort(key=lambda item: item[0], reverse=True)
            top_ratio, top_rect, top_page_index = page_ratios[0]
            runner_up = page_ratios[1][0] if len(page_ratios) > 1 else 0.0
            if top_ratio - runner_up < 0.2:
                continue
            candidate_weight = float(candidate.get("weight", 0.5))
            score = top_ratio * 60 * candidate_weight
            if not best_match or score > best_match["score"]:
                best_match = {
                    "page": top_page_index.page_number,
                    "bbox": _rect_to_bbox(top_rect, _page_index_rect(top_page_index)),
                    "score": score,
                    "locator_strategy": "pymupdf_fuzzy",
                    "matched_text": candidate_text,
                }

    return best_match


def _normalize_reference_pages(root: dict[str, Any]) -> bool:
    raw_pages: list[int] = []
    for _, references in _reference_iter(root):
        for reference in references:
            if isinstance(reference, dict):
                raw_page = _safe_page_number(reference.get("page"))
                if raw_page is None:
                    raw_page = _safe_page_number(reference.get("page_number"))
                if raw_page is not None:
                    raw_pages.append(raw_page)
            elif isinstance(reference, str):
                match = re.search(r"(?:стр\.?|page|p\.)\s*(\d+)", reference, flags=re.IGNORECASE)
                if match:
                    raw_pages.append(int(match.group(1)))

    zero_based = any(page == 0 for page in raw_pages)
    if not zero_based:
        return False

    for _, references in _reference_iter(root):
        for reference in references:
            if not isinstance(reference, dict):
                continue
            raw_page = _safe_page_number(reference.get("page"))
            if raw_page is not None:
                reference["page"] = raw_page + 1
            raw_page_number = _safe_page_number(reference.get("page_number"))
            if raw_page_number is not None:
                reference["page_number"] = raw_page_number + 1
    return True


def _enrich_references_with_pdf_geometry(
    extracted_data: dict[str, Any],
    *,
    local_path: str,
    content_type: str | None,
) -> dict[str, Any]:
    metadata = {
        "enabled": GEOMETRY_ENRICHMENT_ENABLED,
        "provider": "pymupdf" if PYMUPDF_INSTALLED else None,
        "applied": False,
        "searchable_pdf": False,
        "ocr_applied": False,
        "page_count": 1,
        "reference_count": 0,
        "matched_reference_count": 0,
        "converted_string_reference_count": 0,
        "synthetic_reference_count": 0,
        "zero_based_normalized": False,
        "errors": [],
    }

    if not GEOMETRY_ENRICHMENT_ENABLED:
        return metadata
    if not PYMUPDF_INSTALLED:
        metadata["errors"].append("PyMuPDF is not installed")
        return metadata
    if (content_type or "").split(";", 1)[0].strip().lower() != "application/pdf":
        metadata["errors"].append("Geometry enrichment currently supports PDF only")
        return metadata

    document = None
    _tables_cache.clear()
    try:
        document, pages = _build_pdf_page_index(local_path)
        metadata["page_count"] = document.page_count
        # Текстовый слой годен, только если слова осмысленные. У PDF с битой
        # кодировкой шрифта слова «есть», но это мусор — по нему цитаты не
        # находятся (0 координат, клик не работает). Тогда падаем на OCR.
        has_words = any(page.words for page in pages)
        # Достаточно ли слов: скан с парой слов на странице («осмысленных», но
        # пустых по сути) тоже надо отправить на OCR, иначе цитаты не находятся.
        enough_words = _page_words_are_enough(pages)
        text_layer_usable = has_words and _page_words_are_usable(pages) and enough_words
        if has_words and not text_layer_usable:
            if not enough_words:
                logger.info("Geometry: text layer too sparse (scanned) — falling back to OCR index")
            else:
                metadata["garbled_text_layer"] = True
                logger.info("Geometry: text layer garbled — falling back to OCR index")
        metadata["searchable_pdf"] = text_layer_usable
        if not metadata["searchable_pdf"]:
            # Attempt OCR fallback for scanned PDFs OR garbled text layer
            ocr_available = False
            try:
                import pytesseract
                pytesseract.get_tesseract_version()
                ocr_available = True
            except Exception:
                pass

            if not ocr_available:
                metadata["errors"].append("PDF has no searchable text layer and tesseract is not available")
                return metadata

            try:
                if document is not None:
                    document.close()
                document, pages = _build_pdf_page_index_ocr(local_path)
                metadata["searchable_pdf"] = any(page.words for page in pages)
                metadata["ocr_applied"] = True
            except Exception as exc:
                logger.warning("OCR fallback failed: %s", exc)
                metadata["errors"].append(f"OCR failed: {exc}")
                return metadata

            if not metadata["searchable_pdf"]:
                metadata["errors"].append("OCR produced no words")
                return metadata

        metadata["zero_based_normalized"] = _normalize_reference_pages(extracted_data)
        metadata["synthetic_reference_count"] = _inject_synthetic_references(extracted_data)
        generic_anchors = _collect_generic_anchor_texts(extracted_data)

        for holder, references, ancestors in _reference_iter_with_ancestors(extracted_data):
            label_context = _collect_label_context(holder)
            value_context = _collect_value_context(holder)
            ancestor_context = _collect_ancestor_context(ancestors)
            for index, reference in enumerate(references):
                metadata["reference_count"] += 1
                original_reference = reference
                if isinstance(reference, str):
                    reference = {
                        "quote_text": reference,
                        "anchor_text": reference,
                        "locator_text": reference,
                    }
                    references[index] = reference
                    metadata["converted_string_reference_count"] += 1
                if not isinstance(reference, dict):
                    continue

                candidates = _collect_reference_candidates(
                    reference,
                    label_context=label_context,
                    value_context=value_context,
                    generic_anchors=generic_anchors,
                )
                if not candidates:
                    continue

                match = _find_reference_location(pages, candidates, ancestor_context + label_context)
                if not match:
                    # Геометрия не смогла привязать цитату к месту в PDF.
                    # Для сканов (OCR) номер страницы от LLM — это догадка
                    # (обычно дефолтная "1"), показывать её как точную позицию нельзя:
                    # это и приводило к тому, что все характеристики падали на стр. 1.
                    # Помечаем референс как непроверенный и убираем недостоверную страницу.
                    if metadata["ocr_applied"]:
                        reference["page"] = None
                        reference["page_number"] = None
                    reference["position_unverified"] = True
                    if isinstance(original_reference, dict):
                        original_reference.update(reference)
                    continue

                reference["page"] = match["page"]
                reference["page_number"] = match["page"]
                reference["bbox"] = match["bbox"]
                reference["locator_strategy"] = match["locator_strategy"]
                reference.setdefault("quote_text", candidates[0]["text"])
                reference.setdefault("anchor_text", candidates[0]["text"])
                reference["geometry_source"] = "pymupdf"
                reference["position_unverified"] = False
                # matched_text сохраняем только если он согласуется с quote_text.
                # В таблицах геометрия часто находит строку по НАЗВАНИЮ характеристики
                # ("Максимальный напор, м"), тогда как quote_text — это ЗНАЧЕНИЕ ("35").
                # Это валидная привязка (нашли нужную строку, где значение и стоит),
                # но фронтенд считает matched!=quote «обманчивым якорем» и прячет такие
                # совпадения — из-за этого в паспорте отображалась лишь 1 характеристика.
                # Чтобы не терять корректные bbox, в таком случае matched_text не выставляем.
                matched_text = match.get("matched_text")
                quote_text = reference.get("quote_text")
                if _matched_text_consistent_with_quote(matched_text, quote_text):
                    reference["matched_text"] = matched_text
                else:
                    reference.pop("matched_text", None)
                reference["match_score"] = round(float(match.get("score", 0.0)), 3)
                metadata["matched_reference_count"] += 1

                if isinstance(original_reference, dict):
                    original_reference.update(reference)

        metadata["applied"] = metadata["matched_reference_count"] > 0
        return metadata
    except Exception as exc:
        logger.exception("PyMuPDF geometry enrichment failed")
        metadata["errors"].append(str(exc))
        return metadata
    finally:
        if document is not None:
            try:
                document.close()
            except Exception:
                logger.warning("Failed to close PDF document during geometry enrichment")


def _get_downloaded_file_bytes(downloaded_file: DownloadedFile) -> bytes:
    if downloaded_file.file_bytes is None:
        downloaded_file.file_bytes = Path(downloaded_file.local_path).read_bytes()
    return downloaded_file.file_bytes


def _cleanup_downloaded_file(downloaded_file: DownloadedFile) -> None:
    if not downloaded_file.local_path:
        return
    try:
        os.remove(downloaded_file.local_path)
    except OSError:
        logger.warning("Failed to remove temp file %s", downloaded_file.local_path)


def _build_result_page(
    extracted_data: dict[str, Any],
    *,
    raw_text: Optional[str] = None,
    page_no: int = 1,
) -> dict[str, Any]:
    return {
        "page_no": page_no,
        "extracted_data": extracted_data,
        "raw_text": raw_text,
        "errors": None,
    }


def _build_openai_headers(api_key: Optional[str]) -> dict[str, str]:
    headers = {"Content-Type": "application/json"}
    if api_key:
        headers["Authorization"] = f"Bearer {api_key}"
    return headers


async def _chat_completion_json(
    *,
    endpoint: str,
    headers: dict[str, str],
    payload: dict[str, Any],
    provider_name: str,
    repair_schema: Optional[dict[str, Any]] = None,
    repair_model: Optional[str] = None,
) -> tuple[dict[str, Any], dict[str, Any], bool]:
    used_fallback = False
    model_key = str(payload.get("model") or "")
    skip_strict = model_key and _strict_schema_supported.get(model_key) is False

    actual_payload = payload
    if skip_strict and "response_format" in payload:
        actual_payload = dict(payload)
        actual_payload.pop("response_format", None)
        if isinstance(actual_payload.get("provider"), dict):
            provider = dict(actual_payload["provider"])
            provider.pop("require_parameters", None)
            actual_payload["provider"] = provider if provider else actual_payload.pop("provider", None)
        used_fallback = True

    async with httpx.AsyncClient(
        timeout=REMOTE_API_TIMEOUT_SECONDS,
        trust_env=False,
    ) as client:
        response = await client.post(endpoint, json=actual_payload, headers=headers)

        if response.status_code >= 400 and "response_format" in actual_payload:
            if model_key:
                _strict_schema_supported[model_key] = False
            fallback_payload = dict(actual_payload)
            fallback_payload.pop("response_format", None)
            if isinstance(fallback_payload.get("provider"), dict):
                provider = dict(fallback_payload["provider"])
                provider.pop("require_parameters", None)
                if provider:
                    fallback_payload["provider"] = provider
                else:
                    fallback_payload.pop("provider", None)
            used_fallback = True
            retry_response = await client.post(endpoint, json=fallback_payload, headers=headers)
            if retry_response.status_code >= 400:
                _raise_provider_http_error(retry_response, provider_name)
            data = retry_response.json()
        else:
            if response.status_code >= 400:
                _raise_provider_http_error(response, provider_name)
            if model_key and not skip_strict and "response_format" in payload:
                _strict_schema_supported[model_key] = True
            data = response.json()

    if not isinstance(data, dict):
        raise RuntimeError(f"{provider_name} response is not a JSON object")
    raw_content = _extract_message_json_text(data, provider_name=provider_name)
    try:
        parsed = json.loads(raw_content)
    except json.JSONDecodeError as exc:
        diagnostic = _json_decode_diagnostic(provider_name, raw_content, exc, data)
        if repair_schema and repair_model:
            logger.warning(
                "%s; attempting repair",
                diagnostic,
            )
            try:
                parsed = await _repair_raw_text_to_schema(
                    endpoint=endpoint,
                    headers=headers,
                    model=repair_model,
                    schema=repair_schema,
                    raw_text=raw_content,
                    provider_name=f"{provider_name} raw JSON repair",
                )
            except Exception as repair_exc:
                raise RuntimeError(
                    f"{provider_name} returned invalid JSON and repair failed. "
                    f"Original response diagnostic: {diagnostic}; "
                    f"repair_error={repair_exc}"
                ) from repair_exc
        else:
            raise RuntimeError(diagnostic) from exc
    if not isinstance(parsed, dict):
        # Некоторые модели (gpt-4.1 в fallback без схемы) возвращают JSON-массив
        # на верхнем уровне вместо объекта {products: [...]}. Не падаем, а
        # оборачиваем: список изделий или плоский список характеристик — это
        # валидные данные, дальше их нормализует gateway.
        if isinstance(parsed, list):
            logger.info(
                "%s returned a top-level JSON array — wrapping into {products: [...]}",
                provider_name,
            )
            return {"products": parsed}, data, used_fallback
        raise RuntimeError(f"{provider_name} returned non-object JSON")
    return parsed, data, used_fallback


def _schema_required_keys(schema: dict[str, Any]) -> set[str]:
    required = schema.get("required")
    if not isinstance(required, list):
        return set()
    return {item for item in required if isinstance(item, str)}


def _needs_schema_repair(candidate: dict[str, Any], schema: dict[str, Any]) -> bool:
    required = _schema_required_keys(schema)
    if required and not required.issubset(candidate.keys()):
        return True
    return False


async def _repair_json_to_schema(
    *,
    endpoint: str,
    headers: dict[str, str],
    model: str,
    schema: dict[str, Any],
    candidate: dict[str, Any],
    provider_name: str,
) -> dict[str, Any]:
    repair_prompt = (
        "Преобразуй исходный JSON к целевой схеме. "
        "Сохрани все фактические значения. Не придумывай данные. "
        "Если значения нет, используй null, пустую строку или пустой массив в зависимости от схемы. "
        "Верни только валидный JSON без markdown.\n\n"
        f"Целевая схема:\n{json.dumps(schema, ensure_ascii=False)}\n\n"
        f"Исходный JSON:\n{json.dumps(candidate, ensure_ascii=False)}"
    )
    repair_payload = {
        "model": model,
        "messages": [{"role": "user", "content": repair_prompt}],
        "temperature": 0,
        "max_tokens": OPENROUTER_MAX_TOKENS,
        "stream": False,
        "response_format": _json_schema_response_format("json_schema_repair", schema),
    }
    repaired, _, _ = await _chat_completion_json(
        endpoint=endpoint,
        headers=headers,
        payload=repair_payload,
        provider_name=provider_name,
    )
    return repaired


async def _repair_raw_text_to_schema(
    *,
    endpoint: str,
    headers: dict[str, str],
    model: str,
    schema: dict[str, Any],
    raw_text: str,
    provider_name: str,
) -> dict[str, Any]:
    repair_prompt = (
        "Ниже невалидный или незавершенный JSON, полученный из document extraction. "
        "Преобразуй его в валидный JSON строго по целевой схеме. "
        "Сохрани все фактические данные. Не придумывай новые значения. "
        "Если значения нет, используй null, пустую строку или пустой массив по смыслу схемы. "
        "Верни только валидный JSON без markdown.\n\n"
        f"Целевая схема:\n{json.dumps(schema, ensure_ascii=False)}\n\n"
        f"Сырой ответ:\n{raw_text}"
    )
    repair_payload = {
        "model": model,
        "messages": [{"role": "user", "content": repair_prompt}],
        "temperature": 0,
        "max_tokens": OPENROUTER_MAX_TOKENS,
        "stream": False,
        "response_format": _json_schema_response_format("raw_json_repair", schema),
    }
    repaired, _, _ = await _chat_completion_json(
        endpoint=endpoint,
        headers=headers,
        payload=repair_payload,
        provider_name=provider_name,
    )
    return repaired


def _build_openrouter_messages(
    *,
    prompt: str,
    downloaded_file: DownloadedFile,
) -> list[dict[str, Any]]:
    base64_payload = base64.b64encode(_get_downloaded_file_bytes(downloaded_file)).decode("ascii")
    message_content: list[dict[str, Any]] = [{"type": "text", "text": prompt}]

    if _looks_like_image(downloaded_file.filename, downloaded_file.content_type):
        image_mime = downloaded_file.content_type or "image/jpeg"
        message_content.append(
            {
                "type": "image_url",
                "image_url": {"url": f"data:{image_mime};base64,{base64_payload}"},
            }
        )
    else:
        file_mime = downloaded_file.content_type or "application/pdf"
        message_content.append(
            {
                "type": "file",
                "file": {
                    "filename": downloaded_file.filename,
                    "file_data": f"data:{file_mime};base64,{base64_payload}",
                },
            }
        )

    return [{"role": "user", "content": message_content}]


def _iter_docx_blocks(document: Any):
    body = document.element.body
    for child in body.iterchildren():
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "p":
            yield WordParagraph(child, document)
        elif tag == "tbl":
            yield WordTable(child, document)


def _clean_docx_text(value: str | None) -> str:
    if not value:
        return ""
    return re.sub(r"\s+", " ", value).strip()


def _convert_docx_to_structured_text(local_path: str) -> dict[str, Any]:
    document = WordDocument(local_path)
    lines: list[str] = []
    paragraph_count = 0
    table_count = 0

    for block in _iter_docx_blocks(document):
        if isinstance(block, WordParagraph):
            text = _clean_docx_text(block.text)
            if not text:
                continue
            paragraph_count += 1
            style_name = ""
            try:
                style_name = _clean_docx_text(block.style.name)
            except Exception:
                style_name = ""
            prefix = f"[P{paragraph_count}]"
            if style_name:
                prefix += f" ({style_name})"
            lines.append(f"{prefix} {text}")
            continue

        if isinstance(block, WordTable):
            table_count += 1
            lines.append(f"[T{table_count}] TABLE START")
            seen_rows: set[tuple[str, ...]] = set()
            for row_index, row in enumerate(block.rows, start=1):
                cells = tuple(_clean_docx_text(cell.text) for cell in row.cells)
                if not any(cells):
                    continue
                if cells in seen_rows:
                    continue
                seen_rows.add(cells)
                serialized_cells = " | ".join(cell if cell else "—" for cell in cells)
                lines.append(f"[T{table_count}R{row_index}] {serialized_cells}")
            lines.append(f"[T{table_count}] TABLE END")

    structured_text = "\n".join(lines).strip()
    return {
        "text": structured_text,
        "paragraph_count": paragraph_count,
        "table_count": table_count,
    }


def _convert_xlsx_to_structured_text(local_path: str) -> dict[str, Any]:
    """Извлекает текст из Excel-книги для передачи в LLM.

    Каждый лист сериализуется как таблица: строки с непустыми ячейками,
    разделённые ' | '. Формат повторяет DOCX-таблицы ('[T..R..] ...'), чтобы
    LLM одинаково понимал табличную структуру независимо от исходного формата.
    Пустые строки и полностью пустые листы пропускаются."""
    import openpyxl

    workbook = openpyxl.load_workbook(local_path, read_only=True, data_only=True)
    lines: list[str] = []
    sheet_count = 0
    row_count = 0
    try:
        for sheet in workbook.worksheets:
            rows_serialized: list[str] = []
            seen_rows: set[tuple[str, ...]] = set()
            for row in sheet.iter_rows(values_only=True):
                cells = tuple(
                    re.sub(r"\s+", " ", str(value)).strip() if value is not None else ""
                    for value in row
                )
                # Обрезаем хвост пустых ячеек, чтобы '—' не плодились до конца листа.
                while cells and cells[-1] == "":
                    cells = cells[:-1]
                if not any(cells):
                    continue
                if cells in seen_rows:
                    continue
                seen_rows.add(cells)
                rows_serialized.append(cells)
            if not rows_serialized:
                continue
            sheet_count += 1
            sheet_name = _clean_docx_text(sheet.title) or f"Sheet{sheet_count}"
            lines.append(f"[S{sheet_count}: {sheet_name}] SHEET START")
            for row_index, cells in enumerate(rows_serialized, start=1):
                row_count += 1
                serialized = " | ".join(cell if cell else "—" for cell in cells)
                lines.append(f"[S{sheet_count}R{row_index}] {serialized}")
            lines.append(f"[S{sheet_count}] SHEET END")
    finally:
        workbook.close()

    structured_text = "\n".join(lines).strip()
    return {
        "text": structured_text,
        "sheet_count": sheet_count,
        "row_count": row_count,
    }


def _convert_office_document_to_pdf(local_path: str, output_dir: str) -> str:
    source_path = Path(local_path)
    target_dir = Path(output_dir)
    target_dir.mkdir(parents=True, exist_ok=True)
    profile_dir = target_dir / "lo-profile"
    profile_dir.mkdir(parents=True, exist_ok=True)
    command = [
        "soffice",
        "--headless",
        "--nologo",
        "--nofirststartwizard",
        "--nodefault",
        f"-env:UserInstallation=file://{profile_dir.as_posix()}",
        "--convert-to",
        "pdf",
        "--outdir",
        str(target_dir),
        str(source_path),
    ]
    # Изолируем HOME/GNUPGHOME в temp-профиль: иначе LibreOffice трогает общий
    # gpg-agent и плодит фоновые gpg-процессы, которые становятся зомби.
    env = dict(os.environ)
    env["HOME"] = str(profile_dir)
    env["GNUPGHOME"] = str(profile_dir / "gnupg")
    result = subprocess.run(
        command, capture_output=True, text=True, timeout=90, check=False, env=env
    )
    if result.returncode != 0:
        detail = (result.stderr or result.stdout or "LibreOffice conversion failed").strip()
        raise RuntimeError(detail)

    expected_path = target_dir / f"{source_path.stem}.pdf"
    if expected_path.exists():
        return str(expected_path)

    candidates = sorted(target_dir.glob("*.pdf"), key=lambda path: path.stat().st_mtime, reverse=True)
    if not candidates:
        raise RuntimeError("LibreOffice did not create a PDF preview")
    return str(candidates[0])


def _build_openrouter_messages_from_text(
    *,
    prompt: str,
    document_text: str,
    filename: str,
) -> list[dict[str, Any]]:
    return [
        {
            "role": "user",
            "content": (
                f"{prompt}\n\n"
                "Ниже содержимое документа, предварительно извлеченное локально. "
                "Сохраняй структуру документа, учитывай параграфы, таблицы и маркеры блоков. "
                "Если возвращаешь references, опирайся на точные фрагменты текста или строки таблиц из этого представления. "
                "quote_text должен быть дословной цитатой из представления, без пересказа. "
                "Если в представлении есть маркеры [PAGE N], указывай page именно по ним. "
                "Если точная цитата не найдена, используй ближайший дословный фрагмент и anchor_text.\n\n"
                f"Имя файла: {filename}\n"
                "Формат представления:\n"
                "- [PAGE N] — номер страницы исходного документа\n"
                "- [P<N>] — параграф\n"
                "- [T<N>] — таблица\n"
                "- [T<N>R<M>] — строка таблицы\n\n"
                f"Содержимое документа:\n{document_text}"
            ),
        }
    ]


async def _llamaparse_upload(client: httpx.AsyncClient, local_path: str, filename: str) -> str:
    url = f"{LLAMAPARSE_BASE_URL.rstrip('/')}/api/parsing/upload"
    headers = {"Authorization": f"Bearer {LLAMAPARSE_API_KEY}"}
    upload_mime = _normalized_content_type(filename, None) or "application/octet-stream"
    with open(local_path, "rb") as fh:
        files = {"file": (filename, fh, upload_mime)}
        data = {"language": LLAMAPARSE_LANGUAGE}
        response = await _llamaparse_request(
            client,
            "POST",
            url,
            headers=headers,
            files=files,
            data=data,
            operation_name="LlamaParse upload",
        )
    job_id = response.json().get("id")
    if not job_id:
        raise RuntimeError("LlamaParse upload returned no job id")
    return job_id


async def _llamaparse_request(
    client: httpx.AsyncClient,
    method: str,
    url: str,
    *,
    operation_name: str,
    **kwargs: Any,
) -> httpx.Response:
    last_error: Exception | None = None
    retryable_statuses = {408, 409, 425, 429, 500, 502, 503, 504}
    attempts = max(1, LLAMAPARSE_MAX_RETRIES + 1)

    for attempt in range(1, attempts + 1):
        try:
            response = await client.request(method, url, **kwargs)
            if response.is_success:
                return response
            if response.status_code in retryable_statuses and attempt < attempts:
                await asyncio.sleep(min(2 ** (attempt - 1), 6))
                continue
            _raise_provider_http_error(response, operation_name)
        except httpx.HTTPError as exc:
            last_error = exc
            if attempt >= attempts:
                break
            await asyncio.sleep(min(2 ** (attempt - 1), 6))

    raise RuntimeError(f"{operation_name} failed after {attempts} attempts: {last_error}")


async def _llamaparse_wait(client: httpx.AsyncClient, job_id: str) -> None:
    url = f"{LLAMAPARSE_BASE_URL.rstrip('/')}/api/parsing/job/{job_id}"
    headers = {"Authorization": f"Bearer {LLAMAPARSE_API_KEY}"}
    deadline = time.monotonic() + LLAMAPARSE_MAX_WAIT_SECONDS
    while True:
        response = await _llamaparse_request(
            client,
            "GET",
            url,
            headers=headers,
            operation_name="LlamaParse status",
        )
        status = str(response.json().get("status", "")).upper()
        if status == "SUCCESS":
            return
        if status in {"ERROR", "FAILED", "CANCELED"}:
            raise RuntimeError(f"LlamaParse job {job_id} failed: {response.json()}")
        if time.monotonic() > deadline:
            raise RuntimeError(
                f"LlamaParse job {job_id} timed out after {LLAMAPARSE_MAX_WAIT_SECONDS}s"
            )
        await asyncio.sleep(LLAMAPARSE_POLLING_INTERVAL)


def _serialize_llamaparse_result(data: dict[str, Any]) -> tuple[str, dict[str, Any]]:
    pages_payload = data.get("pages")
    page_sections: list[str] = []
    page_count = 0

    if isinstance(pages_payload, list):
        for index, page in enumerate(pages_payload, start=1):
            if not isinstance(page, dict):
                continue
            page_no = (
                _safe_page_number(page.get("page"))
                or _safe_page_number(page.get("page_number"))
                or index
            )
            fragment = (
                page.get("md")
                or page.get("markdown")
                or page.get("text")
                or page.get("content")
                or ""
            )
            if not isinstance(fragment, str) or not fragment.strip():
                continue
            page_sections.append(f"[PAGE {page_no}]\n{fragment.strip()}")
        page_count = len(page_sections)

    top_level_content = (
        data.get(LLAMAPARSE_RESULT_TYPE)
        or data.get("markdown")
        or data.get("text")
        or ""
    )
    if isinstance(top_level_content, str):
        top_level_content = top_level_content.strip()
    else:
        top_level_content = ""

    if not page_sections and top_level_content:
        heuristic_pages = [
            chunk.strip()
            for chunk in re.split(r"(?:\n|^)\s*---+\s*(?:\n|$)", top_level_content)
            if chunk.strip()
        ]
        if len(heuristic_pages) > 1:
            page_sections = [
                f"[PAGE {index}]\n{chunk}"
                for index, chunk in enumerate(heuristic_pages, start=1)
            ]
            page_count = len(page_sections)

    if page_sections:
        content = "\n\n".join(page_sections).strip()
    else:
        content = top_level_content

    if not content:
        raise RuntimeError("LlamaParse returned empty document content")

    metadata = {
        "result_type": LLAMAPARSE_RESULT_TYPE,
        "page_count": page_count or None,
        "has_page_markers": bool(page_sections),
    }
    return content, metadata


async def _llamaparse_get_markdown(client: httpx.AsyncClient, job_id: str) -> tuple[str, dict[str, Any]]:
    url = f"{LLAMAPARSE_BASE_URL.rstrip('/')}/api/parsing/job/{job_id}/result/{LLAMAPARSE_RESULT_TYPE}"
    headers = {"Authorization": f"Bearer {LLAMAPARSE_API_KEY}"}
    response = await _llamaparse_request(
        client,
        "GET",
        url,
        headers=headers,
        operation_name="LlamaParse result",
    )
    return _serialize_llamaparse_result(response.json())


async def _extract_via_llamaparse(payload: ExtractionRequest) -> dict[str, Any]:
    if not LLAMAPARSE_API_KEY:
        raise HTTPException(status_code=500, detail="LLAMAPARSE_API_KEY is not set")
    if not OPENROUTER_API_KEY:
        raise HTTPException(
            status_code=500,
            detail="OPENROUTER_API_KEY is not set (required for structured extraction after LlamaParse parsing)",
        )
    if not payload.prompt and not payload.schema_payload:
        raise HTTPException(status_code=400, detail="schema or prompt is required")

    schema = normalize_json_schema(payload.schema_payload, payload.prompt)
    prompt = payload.prompt or "Extract structured data from the document and return JSON."
    prompt = (
        f"{prompt}\n"
        "Верни результат строго как JSON schema response_format. "
        "Не добавляй markdown, пояснения или кодовые блоки."
    )

    downloaded_file = await _download_file(payload.file_url)
    is_docx = _looks_like_docx(downloaded_file.filename, downloaded_file.content_type)
    docx_pdf_temp_dir: tempfile.TemporaryDirectory | None = None
    started_at = time.monotonic()
    markdown_text = ""
    parse_metadata: dict[str, Any] = {
        "result_type": LLAMAPARSE_RESULT_TYPE,
        "page_count": None,
        "has_page_markers": False,
    }
    geometry_metadata: dict[str, Any] = {
        "enabled": GEOMETRY_ENRICHMENT_ENABLED,
        "provider": "pymupdf" if PYMUPDF_INSTALLED else None,
        "applied": False,
        "searchable_pdf": False,
        "page_count": 1,
        "reference_count": 0,
        "matched_reference_count": 0,
        "converted_string_reference_count": 0,
        "synthetic_reference_count": 0,
        "zero_based_normalized": False,
        "errors": [],
    }
    geometry_local_path = downloaded_file.local_path
    geometry_content_type = downloaded_file.content_type
    try:
        async with httpx.AsyncClient(
            timeout=httpx.Timeout(LLAMAPARSE_REQUEST_TIMEOUT_SECONDS),
            trust_env=False,
        ) as client:
            job_id = await _llamaparse_upload(client, downloaded_file.local_path, downloaded_file.filename)
            logger.info("LlamaParse job created: %s", job_id)
            await _llamaparse_wait(client, job_id)
            markdown_text, parse_metadata = await _llamaparse_get_markdown(client, job_id)
            logger.info(
                "LlamaParse parsing finished in %.2fs, content length=%d",
                time.monotonic() - started_at,
                len(markdown_text),
            )
    except HTTPException:
        raise
    except Exception as exc:
        logger.exception("LlamaParse parsing failed")
        raise HTTPException(
            status_code=502,
            detail=(
                f"LlamaParse parsing failed: {exc}\n\n"
                f"Extraction service traceback:\n{traceback.format_exc()}"
            ),
        ) from exc

    messages = _build_openrouter_messages_from_text(
        prompt=prompt,
        document_text=markdown_text,
        filename=downloaded_file.filename,
    )
    payload_json: dict[str, Any] = {
        "model": OPENROUTER_MODEL,
        "messages": messages,
        "temperature": 0,
        "max_tokens": OPENROUTER_MAX_TOKENS,
        "stream": False,
        "response_format": _json_schema_response_format("llamaparse_extraction", schema),
    }
    provider_preferences = _build_openrouter_provider_preferences(require_parameters=True)
    if provider_preferences:
        payload_json["provider"] = provider_preferences

    headers = _build_openai_headers(OPENROUTER_API_KEY)
    if not LLM_IS_YANDEX:
        headers["HTTP-Referer"] = OPENROUTER_SITE_URL or "http://localhost:8005"
        headers["X-Title"] = OPENROUTER_APP_NAME or "extraction-service"

    try:
        extracted_data, provider_response, used_fallback = await _chat_completion_json(
            endpoint=f"{OPENROUTER_BASE_URL.rstrip('/')}/chat/completions",
            headers=headers,
            payload=payload_json,
            provider_name="OpenRouter (llamaparse)",
            repair_schema=schema,
            repair_model=OPENROUTER_MODEL,
        )
        if used_fallback and _needs_schema_repair(extracted_data, schema):
            extracted_data = await _repair_json_to_schema(
                endpoint=f"{OPENROUTER_BASE_URL.rstrip('/')}/chat/completions",
                headers=headers,
                model=OPENROUTER_MODEL,
                schema=schema,
                candidate=extracted_data,
                provider_name="OpenRouter schema repair (llamaparse)",
            )
        if is_docx:
            try:
                docx_pdf_temp_dir = tempfile.TemporaryDirectory()
                geometry_local_path = await to_thread.run_sync(
                    lambda: _convert_office_document_to_pdf(
                        downloaded_file.local_path,
                        docx_pdf_temp_dir.name,
                    )
                )
                geometry_content_type = "application/pdf"
                parse_metadata["pdf_preview"] = {"created": True}
            except Exception as exc:
                logger.exception("DOCX PDF preview conversion failed")
                parse_metadata["pdf_preview"] = {"created": False, "error": str(exc)}
        geometry_metadata = await to_thread.run_sync(
            lambda: _enrich_references_with_pdf_geometry(
                extracted_data,
                local_path=geometry_local_path,
                content_type=geometry_content_type,
            )
        )
    except HTTPException:
        raise
    except Exception as exc:
        logger.exception("OpenRouter structured extraction failed (llamaparse)")
        raise HTTPException(
            status_code=502,
            detail=(
                f"Structured extraction failed after LlamaParse: {exc}\n\n"
                f"Extraction service traceback:\n{traceback.format_exc()}"
            ),
        ) from exc
    finally:
        _cleanup_downloaded_file(downloaded_file)
        if docx_pdf_temp_dir is not None:
            docx_pdf_temp_dir.cleanup()

    elapsed = time.monotonic() - started_at
    logger.info("LlamaParse full pipeline finished in %.2fs", elapsed)

    return jsonable_encoder({
        "analysis_id": payload.analysis_id,
        "file_id": payload.file_id,
        "file_type": payload.file_type,
        "backend": "llamaparse",
        "model_parse": "llamaparse",
        "model_extract": OPENROUTER_MODEL,
        "result": extracted_data,
        "extraction": {"pages": [_build_result_page(extracted_data, raw_text=markdown_text)]},
        "extraction_metadata": {
            "docling_version": None,
            "page_count": parse_metadata.get("page_count") or geometry_metadata.get("page_count"),
            "errors": [],
            "provider": "llamaparse+openrouter",
            "provider_usage": provider_response.get("usage"),
            "geometry": geometry_metadata,
            "llamaparse": parse_metadata,
            "docx_conversion": None,
        },
    })


async def _extract_via_docling_local(payload: ExtractionRequest) -> dict[str, Any]:
    _raise_docling_backend_unavailable("docling_local")


async def _extract_via_docling_remote(payload: ExtractionRequest) -> dict[str, Any]:
    _raise_docling_backend_unavailable("docling_remote")


async def _extract_via_openrouter(payload: ExtractionRequest) -> dict[str, Any]:
    if not payload.prompt and not payload.schema_payload:
        raise HTTPException(status_code=400, detail="schema or prompt is required")
    if not OPENROUTER_API_KEY:
        raise HTTPException(status_code=500, detail="OPENROUTER_API_KEY is not set")

    schema = normalize_json_schema(payload.schema_payload, payload.prompt)
    prompt = payload.prompt or "Extract structured data from the document and return JSON."
    prompt = (
        f"{prompt}\n"
        "Верни результат строго как JSON schema response_format. "
        "Не добавляй markdown, пояснения или кодовые блоки."
    )

    downloaded_file = await _download_file(payload.file_url)
    is_docx = _looks_like_docx(downloaded_file.filename, downloaded_file.content_type)
    is_excel = _looks_like_excel(downloaded_file.filename, downloaded_file.content_type)
    is_pdf = _looks_like_pdf(downloaded_file.filename, downloaded_file.content_type)
    logger.info(
        "File detected: filename=%r content_type=%r is_docx=%s is_excel=%s is_pdf=%s",
        downloaded_file.filename, downloaded_file.content_type, is_docx, is_excel, is_pdf,
    )
    docx_conversion_metadata: dict[str, Any] | None = None
    excel_conversion_metadata: dict[str, Any] | None = None
    pdf_conversion_metadata: dict[str, Any] | None = None
    docx_pdf_temp_dir: tempfile.TemporaryDirectory | None = None

    if is_docx:
        try:
            docx_conversion_metadata = await to_thread.run_sync(
                lambda: _convert_docx_to_structured_text(downloaded_file.local_path)
            )
        except Exception as exc:
            logger.exception("DOCX conversion failed")
            raise HTTPException(
                status_code=502,
                detail=(
                    f"DOCX conversion failed: {exc}\n\n"
                    f"Extraction service traceback:\n{traceback.format_exc()}"
                ),
            ) from exc
    elif is_excel:
        try:
            excel_conversion_metadata = await to_thread.run_sync(
                lambda: _convert_xlsx_to_structured_text(downloaded_file.local_path)
            )
        except Exception as exc:
            logger.exception("Excel conversion failed")
            raise HTTPException(
                status_code=502,
                detail=(
                    f"Excel conversion failed: {exc}\n\n"
                    f"Extraction service traceback:\n{traceback.format_exc()}"
                ),
            ) from exc
    elif is_pdf and PYMUPDF_INSTALLED:
        # Для PDF: сначала пробуем извлечь текст локально (searchable или OCR).
        # Если текст получен — отправляем как текст (надёжнее, чем file-parser плагин).
        # Если текст пустой — падаем обратно на file-parser (нативный путь OpenRouter).
        # Для Yandex: file-parser недоступен, поэтому всегда используем текстовый путь.
        try:
            pdf_conversion_metadata = await to_thread.run_sync(
                lambda: _convert_pdf_to_structured_text(
                    downloaded_file.local_path,
                    target_names=_extract_target_names_from_prompt(prompt),
                )
            )
            logger.info(
                "PDF text extraction: ocr_applied=%s text_len=%d",
                pdf_conversion_metadata.get("ocr_applied"),
                len(pdf_conversion_metadata.get("text") or ""),
            )
        except Exception as exc:
            logger.warning("PDF text extraction failed, falling back to file-parser: %s", exc)
            pdf_conversion_metadata = None

    # Определяем текст документа для text-пути
    use_text_path = False
    document_text = ""
    if is_docx and docx_conversion_metadata:
        use_text_path = True
        document_text = (docx_conversion_metadata or {}).get("text", "")
    elif is_excel and excel_conversion_metadata:
        use_text_path = True
        document_text = (excel_conversion_metadata or {}).get("text", "")
    elif is_pdf and pdf_conversion_metadata:
        extracted_text = pdf_conversion_metadata.get("text", "")
        ocr_was_applied = pdf_conversion_metadata.get("ocr_applied", False)
        if extracted_text.strip() and not ocr_was_applied:
            use_text_path = True
            document_text = extracted_text
        elif extracted_text.strip() and ocr_was_applied and LLM_IS_YANDEX:
            use_text_path = True
            document_text = extracted_text
        elif extracted_text.strip() and ocr_was_applied:
            logger.info("Scanned PDF: OCR text available but preferring file-parser for better table reading")
        elif LLM_IS_YANDEX:
            # Yandex doesn't support file content parts — if text extraction yielded nothing,
            # we still can't send the raw file; raise an informative error.
            raise HTTPException(
                status_code=422,
                detail="Yandex AI: PDF has no extractable text layer and file upload is not supported.",
            )
    elif LLM_IS_YANDEX and _looks_like_image(downloaded_file.filename, downloaded_file.content_type):
        raise HTTPException(
            status_code=422,
            detail="Yandex AI: image extraction is not supported (no vision API).",
        )

    logger.info("Extraction path: use_text_path=%s document_text_len=%d", use_text_path, len(document_text))

    messages = (
        _build_openrouter_messages_from_text(
            prompt=prompt,
            document_text=document_text,
            filename=downloaded_file.filename,
        )
        if use_text_path
        else _build_openrouter_messages(
            prompt=prompt,
            downloaded_file=downloaded_file,
        )
    )
    payload_json: dict[str, Any] = {
        "model": OPENROUTER_MODEL,
        "messages": messages,
        "temperature": 0,
        "max_tokens": OPENROUTER_MAX_TOKENS,
        "stream": False,
        "response_format": _json_schema_response_format("openrouter_extraction", schema),
    }
    provider_preferences = _build_openrouter_provider_preferences(require_parameters=True)
    if provider_preferences:
        payload_json["provider"] = provider_preferences
    # file-parser плагин — только для OpenRouter, когда не используем text-путь
    if not LLM_IS_YANDEX and not use_text_path and not _looks_like_image(downloaded_file.filename, downloaded_file.content_type):
        payload_json["plugins"] = [
            {
                "id": "file-parser",
                "pdf": {"engine": OPENROUTER_PDF_ENGINE},
            }
        ]

    headers = _build_openai_headers(OPENROUTER_API_KEY)
    if not LLM_IS_YANDEX:
        headers["HTTP-Referer"] = OPENROUTER_SITE_URL or "http://localhost:8005"
        headers["X-Title"] = OPENROUTER_APP_NAME or "extraction-service"

    started_at = time.monotonic()
    geometry_metadata: dict[str, Any] = {
        "enabled": GEOMETRY_ENRICHMENT_ENABLED,
        "provider": "pymupdf" if PYMUPDF_INSTALLED else None,
        "applied": False,
        "searchable_pdf": False,
        "page_count": 1,
        "reference_count": 0,
        "matched_reference_count": 0,
        "converted_string_reference_count": 0,
        "zero_based_normalized": False,
        "errors": [],
    }
    geometry_local_path = downloaded_file.local_path
    geometry_content_type = downloaded_file.content_type
    try:
        extracted_data, provider_response, used_fallback = await _chat_completion_json(
            endpoint=f"{OPENROUTER_BASE_URL.rstrip('/')}/chat/completions",
            headers=headers,
            payload=payload_json,
            provider_name="OpenRouter",
            repair_schema=schema,
            repair_model=OPENROUTER_MODEL,
        )
        if used_fallback and _needs_schema_repair(extracted_data, schema):
            extracted_data = await _repair_json_to_schema(
                endpoint=f"{OPENROUTER_BASE_URL.rstrip('/')}/chat/completions",
                headers=headers,
                model=OPENROUTER_MODEL,
                schema=schema,
                candidate=extracted_data,
                provider_name="OpenRouter schema repair",
            )
        # Нормализуем references: некоторые модели возвращают dict вместо list
        _normalize_references_in_place(extracted_data)
        # Office-документы (DOCX/Excel) конвертируем в PDF, чтобы геометрия привязала
        # координаты цитат к страницам PDF-превью (вьювер показывает именно его).
        if is_docx or is_excel:
            office_meta = docx_conversion_metadata if is_docx else excel_conversion_metadata
            doc_kind = "DOCX" if is_docx else "Excel"
            try:
                docx_pdf_temp_dir = tempfile.TemporaryDirectory()
                geometry_local_path = await to_thread.run_sync(
                    lambda: _convert_office_document_to_pdf(
                        downloaded_file.local_path,
                        docx_pdf_temp_dir.name,
                    )
                )
                geometry_content_type = "application/pdf"
                if office_meta is not None:
                    office_meta["pdf_preview"] = {"created": True}
            except Exception as exc:
                logger.exception("%s PDF preview conversion failed", doc_kind)
                if office_meta is not None:
                    office_meta["pdf_preview"] = {
                        "created": False,
                        "error": str(exc),
                    }
        geometry_metadata = await to_thread.run_sync(
            lambda: _enrich_references_with_pdf_geometry(
                extracted_data,
                local_path=geometry_local_path,
                content_type=geometry_content_type,
            )
        )
    except HTTPException:
        raise
    except Exception as exc:
        logger.exception("OpenRouter extraction failed")
        detail = (
            f"OpenRouter extraction failed: {exc}\n\n"
            f"Extraction service traceback:\n{traceback.format_exc()}"
        )
        if not DEBUG_ERRORS and detail:
            detail = str(detail)
        raise HTTPException(status_code=502, detail=detail) from exc
    finally:
        _cleanup_downloaded_file(downloaded_file)
        if docx_pdf_temp_dir is not None:
            docx_pdf_temp_dir.cleanup()
        elapsed = time.monotonic() - started_at
        logger.info("OpenRouter extraction finished in %.2fs", elapsed)

    response = {
        "analysis_id": payload.analysis_id,
        "file_id": payload.file_id,
        "file_type": payload.file_type,
        "backend": "openrouter",
        "model_parse": OPENROUTER_MODEL,
        "model_extract": OPENROUTER_MODEL,
        "result": extracted_data,
        "extraction": {"pages": [_build_result_page(extracted_data)]},
        "extraction_metadata": {
            "docling_version": _get_docling_version(),
            "page_count": geometry_metadata.get("page_count", 1),
            "errors": [],
            "provider": "openrouter",
            "provider_usage": provider_response.get("usage"),
            "geometry": geometry_metadata,
            "docx_conversion": docx_conversion_metadata,
            "excel_conversion": excel_conversion_metadata,
            "pdf_conversion": pdf_conversion_metadata,
        },
    }
    return jsonable_encoder(response)


RENDER_CACHE_DIR = Path(tempfile.gettempdir()) / "render_pdf_cache"
# Блокировки на ключ кэша: пока один запрос растеризует документ, остальные
# запросы того же документа ждут результат, а не запускают свою растеризацию.
_render_locks: dict[str, asyncio.Lock] = {}
_render_locks_guard = asyncio.Lock()


def _render_cache_key(url: str) -> str:
    """Ключ кэша по СТАБИЛЬНОЙ части URL (путь к объекту в S3), без подписи.
    Presigned-URL для одного файла каждый раз разный (меняется signature), но
    путь к объекту постоянен — кэшируем по нему."""
    path = unquote(urlparse(url).path)
    return hashlib.sha256(path.encode("utf-8")).hexdigest()


async def _get_render_lock(key: str) -> asyncio.Lock:
    async with _render_locks_guard:
        lock = _render_locks.get(key)
        if lock is None:
            lock = asyncio.Lock()
            _render_locks[key] = lock
        return lock


_classify_cache: dict[str, dict[str, bool]] = {}


def _file_content_hash(local_path: str) -> str:
    """Быстрый хеш содержимого файла для кэша классификации: размер + первые/
    последние 256 КБ (достаточно для уникальности PDF без чтения всего файла)."""
    h = hashlib.sha256()
    size = os.path.getsize(local_path)
    h.update(str(size).encode())
    with open(local_path, "rb") as f:
        h.update(f.read(262144))
        if size > 524288:
            f.seek(-262144, os.SEEK_END)
            h.update(f.read(262144))
    return h.hexdigest()


def _classify_pdf_for_preview(local_path: str) -> dict[str, bool]:
    """Классифицирует PDF для превью: что с ним не так и нужна ли растеризация.

    Вердикт кэшируется по хешу содержимого: один и тот же файл, загруженный
    повторно (под другим storage_path/URL), не проходит дорогую классификацию
    (OSD скан-страниц) заново."""
    try:
        content_key = _file_content_hash(local_path)
        cached = _classify_cache.get(content_key)
        if cached is not None:
            return cached
    except Exception:
        content_key = None
    result = _classify_pdf_for_preview_uncached(local_path)
    if content_key is not None:
        # Ограничиваем рост кэша (простая защита от утечки памяти).
        if len(_classify_cache) > 512:
            _classify_cache.clear()
        _classify_cache[content_key] = result
    return result


def _classify_pdf_for_preview_uncached(local_path: str) -> dict[str, bool]:
    """Классифицирует PDF для превью: что с ним не так и нужна ли растеризация.

    Растеризуем (с выпрямлением), если документ показался бы криво или подсветка
    не совпала бы с координатами:
      - rotated:  есть страницы с поворотом (/Rotate != 0) → вьювер покажет боком;
      - garbled:  текстовый слой — мусор (битая кодировка шрифта);
      - non_embedded_fonts: есть шрифты, которые PDF.js не умеет рендерить.
    Чистый PDF (только web-safe шрифты, без поворота, читаемый текст) —
    отдаём оригинал как есть (быстро, постранично, текст выделяется)."""
    info = {
        "rotated": False,
        "garbled": False,
        "non_embedded_fonts": False,
        "scan_rotated": False,
        "needs_rasterization": False,
    }
    # Шрифты, которые PDF.js умеет рендерить сам без растеризации:
    # стандартные 14 PDF-шрифтов + распространённые web-safe (latin-only).
    # Любой другой встроенный шрифт (особенно кириллические TrueType вроде
    # FuturisC, Helios, PragmaticaC и т.п.) PDF.js не отрисует — нужна растеризация.
    pdfjs_safe_fonts = {
        "helvetica", "courier", "times", "symbol", "zapfdingbats", "arial",
        "georgia", "verdana", "tahoma", "trebuchet", "impact",
        "comic", "palatino", "garamond", "bookman", "avantgarde",
    }
    try:
        doc = fitz.open(local_path)
        try:
            pages_to_check = min(doc.page_count, 15)
            text_parts: list[str] = []
            scan_pages: list[int] = []
            for i in range(pages_to_check):
                page = doc.load_page(i)
                if page.rotation:
                    info["rotated"] = True
                page_text = page.get_text("text", sort=True)
                text_parts.append(page_text)
                # Страница-скан (нет текстового слоя) — кандидат на проверку
                # ориентации содержимого: PDF /Rotate=0, но картинка может быть боком.
                if len(page_text.strip()) < 10:
                    scan_pages.append(i)
                for font in doc.get_page_fonts(i):
                    # font = (xref, ext, type, basefont, name, encoding)
                    ext = (font[1] or "").strip().lower()
                    base = (font[3] or "").lower()
                    is_pdfjs_safe = any(s in base for s in pdfjs_safe_fonts)
                    # Растеризуем при любом шрифте вне web-safe списка: PDF.js не
                    # умеет рендерить произвольные встроенные TrueType/Type1 из
                    # PDF-потока — кириллические шрифты показываются квадратиками.
                    if not is_pdfjs_safe:
                        info["non_embedded_fonts"] = True
            combined = "\n".join(t for t in text_parts if t)
            if len(combined) > 50 and not _text_layer_is_usable(combined):
                info["garbled"] = True
            # Скан-страницы без текстового слоя: проверяем ориентацию содержимого
            # через Tesseract OSD. Если контент повёрнут (текст идёт боком) — нужна
            # растеризация с авто-поворотом, иначе превью покажет документ криво и
            # подсветка (привязанная к выпрямленному OCR-индексу) не совпадёт.
            if scan_pages and not info["rotated"]:
                if _scan_pages_are_rotated(doc, scan_pages[:2]):
                    info["scan_rotated"] = True
        finally:
            doc.close()
    except Exception:
        logger.warning("PDF classification failed, defaulting to rasterization", exc_info=True)
        info["needs_rasterization"] = True
        return info

    info["needs_rasterization"] = (
        info["rotated"]
        or info["garbled"]
        or info["non_embedded_fonts"]
        or info["scan_rotated"]
    )
    return info


def _scan_pages_are_rotated(doc: Any, page_indices: list[int]) -> bool:
    """Проверяет через Tesseract OSD, повёрнуто ли содержимое скан-страниц.
    Возвращает True, если хотя бы одна страница уверенно распознана повёрнутой
    (rotate != 0). Используется для превью: PDF /Rotate=0, но картинка боком.

    Оптимизация скорости: рендер при низком DPI (для OSD достаточно ~100), и
    останавливаемся на ПЕРВОЙ повёрнутой странице. OSD — дорогая операция (десятки
    секунд на документ при высоком DPI и нескольких страницах), а для определения
    «повёрнут ли скан» хватает грубого превью одной-двух страниц."""
    try:
        import pytesseract
        from PIL import Image
        import io
    except Exception:
        return False
    # 100 DPI достаточно для OSD и вдвое-втрое быстрее 150; psm=0 — только OSD.
    osd_config = "--psm 0"
    for i in page_indices:
        try:
            page = doc.load_page(i)
            pix = page.get_pixmap(matrix=fitz.Matrix(100 / 72, 100 / 72), alpha=False)
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            osd = pytesseract.image_to_osd(
                img,
                output_type=pytesseract.Output.DICT,
                config=osd_config,
                timeout=OCR_OSD_TIMEOUT_SECONDS,
            )
            rotate = int(osd.get("rotate", 0) or 0)
            conf = float(osd.get("orientation_conf", 0) or 0)
            if rotate % 360 != 0 and conf >= 2.0:
                return True  # нашли повёрнутую — дальше проверять незачем
        except Exception:
            continue
    return False


def _pdf_needs_rasterization(local_path: str) -> bool:
    return _classify_pdf_for_preview(local_path)["needs_rasterization"]


def _rasterize_pdf(local_path: str, auto_orient: bool = False) -> bytes:
    """Растеризует каждую страницу PDF в JPEG и собирает новый PDF.

    Страницы вставляются как JPEG (а не lossless): размер итогового PDF в разы
    меньше, что критично для скорости загрузки превью на медленном интернете.
    DPI и качество JPEG настраиваются через RENDER_PDF_DPI / RENDER_PDF_JPEG_QUALITY.

    auto_orient=True — выпрямляет повёрнутые страницы (через OSD), чтобы вьювер
    показывал документ прямо И в той же ориентации, в которой геометрия привязала
    координаты (иначе подсветка не совпадёт)."""
    import io
    from PIL import Image

    src = fitz.open(local_path)
    out = fitz.open()
    dpi = RENDER_PDF_DPI
    scale = dpi / 72
    mat = fitz.Matrix(scale, scale)
    for page in src:
        pix = page.get_pixmap(matrix=mat, alpha=False)
        img = Image.open(io.BytesIO(pix.tobytes("png")))
        if auto_orient:
            img = _auto_orient_for_ocr(img)
        # Размер страницы в пунктах при 72 dpi (после возможного поворота)
        pw = img.width * 72 / dpi
        ph = img.height * 72 / dpi
        jpeg_buf = io.BytesIO()
        img.convert("RGB").save(jpeg_buf, format="JPEG", quality=RENDER_PDF_JPEG_QUALITY)
        out_page = out.new_page(width=pw, height=ph)
        out_page.insert_image(fitz.Rect(0, 0, pw, ph), stream=jpeg_buf.getvalue())
    src.close()
    buf = io.BytesIO()
    # JPEG уже сжат — deflate на поток картинок не нужен, только чистим объекты.
    out.save(buf, garbage=2)
    out.close()
    return buf.getvalue()


def _render_response(pdf_bytes: bytes) -> Response:
    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={
            "Content-Disposition": "inline; filename=preview.pdf",
            "Cache-Control": "private, max-age=3600",
        },
    )


@app.get("/render-pdf")
async def render_pdf(url: str) -> Response:
    """Скачивает PDF по URL, рендерит каждую страницу через PyMuPDF в изображение
    и возвращает новый PDF из растровых страниц. Решает проблему с нестандартными
    кириллическими шрифтами, которые PDF.js не умеет отображать.

    Результат кэшируется на диске по ключу объекта S3: повторные открытия того же
    документа отдаются из кэша мгновенно, без повторной растеризации (она грузит CPU)."""
    if not PYMUPDF_INSTALLED:
        raise HTTPException(status_code=501, detail="PyMuPDF not available")

    parsed = urlparse(url)
    if parsed.scheme not in {"http", "https"}:
        raise HTTPException(status_code=400, detail="url must be http or https")

    cache_key = _render_cache_key(url)
    cache_path = RENDER_CACHE_DIR / f"{cache_key}.pdf"
    if cache_path.exists():
        return _render_response(cache_path.read_bytes())

    # Сериализуем растеризацию одного и того же документа: параллельные открытия
    # не должны запускать N одновременных растеризаций (это и сатурировало CPU).
    lock = await _get_render_lock(cache_key)
    async with lock:
        if cache_path.exists():
            return _render_response(cache_path.read_bytes())

        try:
            downloaded = await _download_file(url)
        except Exception as exc:
            raise HTTPException(status_code=502, detail=f"Failed to download file: {exc}") from exc

        try:
            # Растеризуем ТОЛЬКО если это реально нужно (есть невстроенные шрифты).
            # Для обычных PDF со встроенными шрифтами отдаём оригинал — PDF.js
            # покажет его сам, постранично и быстро (без тяжёлой растеризации и
            # без раздувания размера, которое тормозило превью на медленном инете).
            needs_raster = await to_thread.run_sync(
                lambda: _pdf_needs_rasterization(downloaded.local_path)
            )
            if needs_raster:
                # Растеризуем с авто-поворотом: вьювер покажет документ прямо и в
                # той же ориентации, в которой геометрия привязала координаты.
                pdf_bytes = await to_thread.run_sync(
                    lambda: _rasterize_pdf(downloaded.local_path, auto_orient=True)
                )
            else:
                pdf_bytes = await to_thread.run_sync(
                    lambda: Path(downloaded.local_path).read_bytes()
                )
        except Exception as exc:
            raise HTTPException(status_code=500, detail=f"PDF rasterization failed: {exc}") from exc
        finally:
            _cleanup_downloaded_file(downloaded)

        try:
            RENDER_CACHE_DIR.mkdir(parents=True, exist_ok=True)
            tmp_path = cache_path.with_suffix(".pdf.tmp")
            tmp_path.write_bytes(pdf_bytes)
            tmp_path.replace(cache_path)
        except Exception as exc:
            logger.warning("Failed to cache rendered PDF: %s", exc)

    return _render_response(pdf_bytes)


@app.post("/extract")
async def extract_document(payload: ExtractionRequest) -> dict:
    if not payload.schema_payload and not payload.prompt:
        raise HTTPException(status_code=400, detail="schema or prompt is required")

    parsed_url = urlparse(payload.file_url)
    if parsed_url.scheme not in {"http", "https"}:
        raise HTTPException(status_code=400, detail="file_url must be a public URL")

    backend = _select_backend(payload.backend)
    if backend == "docling_local":
        return await _extract_via_docling_local(payload)
    if backend == "docling_remote":
        return await _extract_via_docling_remote(payload)
    if backend == "openrouter":
        return await _extract_via_openrouter(payload)
    if backend == "llamaparse":
        return await _extract_via_llamaparse(payload)
    raise HTTPException(status_code=400, detail=f"Unsupported backend: {backend}")
